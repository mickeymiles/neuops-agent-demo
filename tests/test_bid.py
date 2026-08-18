# -*- coding: utf-8 -*-
"""投标工作台测试：项目管理 / 上传 / 拆标 / 生成 / 自检 / 导出 / 聊天联动
# 规格编号: NO-009 FR-1~FR-7 / NFR-1 / TC-2
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

from main import app  # noqa: E402

client = TestClient(app)


@pytest.fixture
def no_llm(monkeypatch):
    """禁用 DeepSeek 调用：拆标/提炼走规则兜底，保证测试稳定快速（FR-13/14/15）"""
    monkeypatch.setattr("app.agent_chat._load_deepseek_key", lambda: "")
    yield

BID_TXT = (
    "一、资质要求：投标人须具备 ISO9001 质量管理体系认证资质，并提供证书复印件。\n"
    "二、业绩要求：投标人近 3 年须具有同类政务云运维项目业绩不少于 2 个，提供合同与验收报告。\n"
    "三、技术参数：CPU 主频不低于 3.0GHz，内存不小于 64GB（★关键参数）。\n"
    "四、评分标准：技术方案完整性 20 分，项目业绩 15 分。\n"
    "五、废标条款：投标文件未按要求密封的，作废标处理。\n"
)


def _mk_project(name="测试投标项目"):
    r = client.post("/api/bid/projects", json={"name": name, "tenderee": "某市采购中心",
                                               "industry": "政务", "budget": 500, "deadline": "2026-09-30"})
    assert r.status_code == 200
    return r.json()["project"]


def _upload(pid, content="", name="规范书.txt"):
    files = {"files": (name, content.encode("utf-8"), "text/plain")}
    return client.post(f"/api/bid/projects/{pid}/upload", files=files)


# ==================== FR-1 项目管理 ====================

def test_bid_project_crud():
    # NO-009 FR-1: 新建
    p = _mk_project("CRUD项目")
    pid = p["id"]
    assert p["status"] == "草稿"
    assert p["tenderee"] == "某市采购中心"
    # 列表
    r = client.get("/api/bid/projects")
    assert any(x["id"] == pid for x in r.json()["projects"])
    # 更新
    r = client.patch(f"/api/bid/projects/{pid}", json={"industry": "金融"})
    assert r.json()["project"]["industry"] == "金融"
    # 详情
    r = client.get(f"/api/bid/projects/{pid}")
    assert r.status_code == 200
    assert r.json()["project"]["id"] == pid
    # 删除
    r = client.delete(f"/api/bid/projects/{pid}")
    assert r.json()["success"] is True
    assert client.get(f"/api/bid/projects/{pid}").status_code == 404


def test_bid_project_name_required():
    # NO-009 FR-1: 名称必填
    r = client.post("/api/bid/projects", json={"name": "  "})
    assert r.status_code == 400


# ==================== FR-2 上传与文本抽取 ====================

def test_bid_upload_txt():
    # NO-009 FR-2: 上传 txt → 原文件保存 + 文本抽取 + 状态流转
    p = _mk_project("上传项目")
    pid = p["id"]
    r = _upload(pid, BID_TXT)
    assert r.json()["success"] is True
    assert len(r.json()["saved"]) == 1
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert detail["status"] == "已上传"
    assert any(f["name"] == "规范书.txt" for f in detail["files"])


def test_bid_upload_unsupported_ext():
    # NO-009 FR-2: 不支持格式应失败
    p = _mk_project("格式项目")
    pid = p["id"]
    files = {"files": ("bad.exe", b"MZ", "application/octet-stream")}
    r = client.post(f"/api/bid/projects/{pid}/upload", files=files)
    assert r.json()["success"] is False
    assert r.json()["failed"]


# ==================== FR-3 拆标解析 ====================

def test_bid_parse_rule_fallback():
    # NO-009 FR-3 / TC-2: 无 LLM 时规则粗筛仍产出六类骨架，缺项标"未识别"
    from app.bidding.bid_engine import rule_parse
    report = rule_parse(BID_TXT)
    for k in ("qualifications", "performance", "tech_params", "scoring",
              "rejection_clauses", "response_checklist"):
        assert k in report
    assert any("ISO9001" in str(x) for x in report["qualifications"])
    assert any("政务云" in str(x) for x in report["performance"])
    assert any("CPU" in str(x) for x in report["tech_params"])
    # 空文本 → 未识别
    empty = rule_parse("")
    assert empty["qualifications"][0]["item"] == "未识别"


def test_bid_parse_api_flow():
    # NO-009 FR-3: 触发拆标 → 报告六类齐全 + 状态"已拆标"
    p = _mk_project("拆标项目")
    pid = p["id"]
    _upload(pid, BID_TXT)
    r = client.post(f"/api/bid/projects/{pid}/parse")
    assert r.status_code == 200
    report = r.json()["report"]
    assert set(report.keys()) >= {"qualifications", "performance", "tech_params",
                                  "scoring", "rejection_clauses", "response_checklist"}
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert detail["status"] == "已拆标"


# ==================== FR-4 生成材料 ====================

def test_bid_generate_docs():
    # NO-009 FR-4: 四类材料可生成，无依据项标"待补充材料"
    p = _mk_project("生成项目")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    for doc_type in ("tech_proposal", "response", "ppt_outline", "impl_plan"):
        r = client.post(f"/api/bid/projects/{pid}/generate", json={"type": doc_type})
        assert r.status_code == 200, doc_type
        doc = r.json()["doc"]
        assert doc["type"] == doc_type
        assert os.path.isfile(doc["path"])
        assert os.path.getsize(doc["path"]) > 0
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert len(detail["generated_docs"]) == 4
    assert detail["status"] == "已生成"


def test_bid_generate_invalid_type():
    # NO-009 FR-4: 非法材料类型拒绝
    p = _mk_project("非法类型")
    pid = p["id"]
    r = client.post(f"/api/bid/projects/{pid}/generate", json={"type": "unknown"})
    assert r.status_code == 400


# ==================== FR-5 合规自检 ====================

def test_bid_check_compliance():
    # NO-009 FR-5: 自检输出未响应项/红线/评分建议
    p = _mk_project("自检项目")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    client.post(f"/api/bid/projects/{pid}/generate", json={"type": "response"})
    r = client.post(f"/api/bid/projects/{pid}/check")
    assert r.status_code == 200
    result = r.json()["result"]
    assert set(result.keys()) == {"unresponded", "redlines", "scoring"}
    assert isinstance(result["unresponded"], list)
    assert isinstance(result["redlines"], list)
    assert isinstance(result["scoring"], list)
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert detail["status"] == "已自检"


# ==================== FR-6 成果导出 ====================

def test_bid_export_md():
    # NO-009 FR-6: 导出 md
    p = _mk_project("导出项目")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    doc = client.post(f"/api/bid/projects/{pid}/generate", json={"type": "ppt_outline"}).json()["doc"]
    r = client.get(f"/api/bid/projects/{pid}/export/{doc['id']}?fmt=md")
    assert r.status_code == 200
    assert "text/markdown" in r.headers.get("content-type", "")
    assert "#" in r.text


def test_bid_export_docx():
    # NO-009 FR-6: 导出 docx
    p = _mk_project("导出docx")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    doc = client.post(f"/api/bid/projects/{pid}/generate", json={"type": "tech_proposal"}).json()["doc"]
    r = client.get(f"/api/bid/projects/{pid}/export/{doc['id']}?fmt=docx")
    assert r.status_code == 200
    assert "wordprocessingml" in r.headers.get("content-type", "")
    assert len(r.content) > 100


# ==================== FR-7 聊天联动 ====================

def test_bid_employee_prompt_guidance():
    # NO-009 FR-7: emp-007 prompt 含工作台跳转指引；其他员工不受影响
    from app.agent_chat import build_employee_prompt
    prompt_007 = build_employee_prompt("emp-007")
    assert "工作台协作边界" in prompt_007
    assert "/bidding" in prompt_007
    prompt_001 = build_employee_prompt("emp-001")
    assert "工作台协作边界" not in prompt_001


# ==================== 页面与知识库 ====================

def test_bidding_page():
    r = client.get("/bidding")
    assert r.status_code == 200
    assert "投标工作台" in r.text


def test_bid_kb_seeded_and_bound():
    # NO-008 delta: 5 类投标知识库已预置且绑定 emp-007
    from app import db as _db
    from app.seed_bid_kb import BID_KB_SEED, EMPLOYEE_ID
    kbs = {kb["name"]: kb for kb in _db.db_list_knowledge_bases()}
    for name in BID_KB_SEED:
        assert name in kbs, f"投标知识库未预置: {name}"
    bound = _db.db_get_employee_kb_ids(EMPLOYEE_ID)
    bound_names = [kb["name"] for kb in _db.db_list_knowledge_bases() if kb["id"] in bound]
    for name in BID_KB_SEED:
        assert name in bound_names, f"{EMPLOYEE_ID} 未绑定: {name}"


# ==================== FR-10 多文件分类解析 ====================

def test_bid_upload_with_category():
    # NO-009 FR-10: 带分类上传 → extracted/<category>/ 独立落盘 + 文件列表带分类标签
    from app.bidding.bid_engine import BID_UPLOAD_ROOT
    p = _mk_project("分类上传")
    pid = p["id"]
    files = {"files": ("技术规范.txt", BID_TXT.encode("utf-8"), "text/plain")}
    data = {"categories": '{"技术规范.txt":"technical"}'}
    r = client.post(f"/api/bid/projects/{pid}/upload", files=files, data=data)
    assert r.json()["success"] is True
    assert r.json()["saved"][0]["category"] == "technical"
    assert os.path.isfile(os.path.join(BID_UPLOAD_ROOT, str(pid),
                                       "extracted", "technical", "技术规范.txt"))
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    f0 = next(f for f in detail["files"] if f["name"] == "技术规范.txt")
    assert f0["category"] == "technical"


def test_bid_auto_category():
    # NO-009 FR-10: 文件名自动识别 + 内容兜底
    from app.bidding.bid_engine import auto_category
    assert auto_category("可行性研究报告.docx") == "feasibility"
    assert auto_category("技术规范书.pdf") == "technical"
    assert auto_category("商务要求.docx") == "commercial"
    assert auto_category("合同协议书.docx") == "contract"
    assert auto_category("格式编制要求.docx") == "format"
    assert auto_category("随便.txt") == "other"
    # 文件名无特征 → 内容兜底（保守关键词）
    assert auto_category("文档1.txt", "格式要求：正文宋体小四，1.5倍行距") == "format"
    assert auto_category("文档2.txt", "技术参数：CPU主频不低于3.0GHz（★关键参数）") == "technical"


def test_bid_parse_category_priority():
    # NO-009 FR-10: 主拆标仅合并正式标书类别（可研隔离）+ 合并文本带【来源文件】头
    from app.bidding.bid_engine import _read_extracted_text
    p = _mk_project("分类优先级")
    pid = p["id"]
    fea = "一、技术参数：CPU主频不低于2.0GHz（可研阶段初步指标）。\n"
    off = "一、技术参数：CPU主频不低于3.0GHz（★关键参数，以正式标书为准）。\n"
    files = [("files", ("可研文件.txt", fea.encode("utf-8"), "text/plain")),
             ("files", ("技术规范.txt", off.encode("utf-8"), "text/plain"))]
    data = {"categories": '{"可研文件.txt":"feasibility","技术规范.txt":"technical"}'}
    r = client.post(f"/api/bid/projects/{pid}/upload", files=files, data=data)
    assert r.json()["success"] is True
    main = _read_extracted_text(pid, categories=("technical", "commercial", "contract", "other"))
    assert "3.0GHz" in main and "2.0GHz" not in main
    assert "【来源文件：技术规范】" in main
    # 仅可研文件 → 主文本为空，回退全量合并，仍可拆标
    p2 = _mk_project("仅可研")
    pid2 = p2["id"]
    files2 = {"files": ("可研文件.txt", fea.encode("utf-8"), "text/plain")}
    data2 = {"categories": '{"可研文件.txt":"feasibility"}'}
    client.post(f"/api/bid/projects/{pid2}/upload", files=files2, data=data2)
    r2 = client.post(f"/api/bid/projects/{pid2}/parse")
    assert r2.status_code == 200
    assert set(r2.json()["report"].keys()) >= {"qualifications", "tech_params"}


# ==================== FR-11/FR-12 投标格式规范 ====================

def test_bid_extract_format_spec():
    # NO-009 FR-11/FR-12: 从格式要求文本提取字体/字号/行距/页边距/页码
    from app.bidding.bid_engine import extract_format_spec
    text = ("格式要求：投标文件正文采用宋体小四号，1.5倍行距；标题用黑体；"
            "页边距：上2.54cm，下2.54cm，左3.17cm，右3.17cm；页脚居中页码。")
    spec = extract_format_spec(text)
    assert spec["fonts"][0] == "宋体"
    assert spec["size_body"] == 12
    assert spec["line_spacing"] == 1.5
    assert spec["margin_top_cm"] == 2.54
    assert spec["margin_right_cm"] == 3.17
    assert spec["page_number"] is True
    # 无格式要求 → 空 dict（由默认基线兜底）
    assert extract_format_spec("无特殊排版要求") == {}


def test_bid_export_docx_styled():
    # NO-009 FR-11/FR-12: docx 导出含排版基线（宋体正文/黑体标题/页边距/封面）
    import io
    import zipfile
    p = _mk_project("排版docx")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    doc = client.post(f"/api/bid/projects/{pid}/generate", json={"type": "tech_proposal"}).json()["doc"]
    r = client.get(f"/api/bid/projects/{pid}/export/{doc['id']}?fmt=docx")
    assert r.status_code == 200
    with zipfile.ZipFile(io.BytesIO(r.content)) as z:
        styles = z.read("word/styles.xml").decode("utf-8", "ignore")
        docxml = z.read("word/document.xml").decode("utf-8", "ignore")
    assert "宋体" in styles and "黑体" in styles
    assert "w:pgMar" in docxml
    assert "投 标 文 件" in docxml  # 封面存在（字距空格）


# ==================== FR-13 投标模板 ====================

def _mk_template_docx(path="/tmp/test-bid-template.docx"):
    """构造含章节标题的模板 docx"""
    from docx import Document
    d = Document()
    d.add_heading("技术方案建议书", 0)
    d.add_heading("一、项目理解与总体方案", 1)
    d.add_paragraph("【模板保留】项目理解占位")
    d.add_heading("二、技术方案要点", 1)
    d.add_paragraph("【模板保留】技术要点占位")
    d.add_heading("三、实施方案概述", 1)
    d.add_paragraph("【模板保留】实施占位")
    d.save(path)
    return path


def test_bid_template_upload_structure(no_llm):
    # NO-009 FR-13: 上传 docx 模板 → 章节树解析 + project.template 元信息
    p = _mk_project("模板上传")
    pid = p["id"]
    tpl_path = _mk_template_docx()
    with open(tpl_path, "rb") as f:
        r = client.post(f"/api/bid/projects/{pid}/template",
                        files={"file": ("投标模板.docx", f.read(),
                                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200, r.text
    tpl = r.json()["template"]
    assert tpl["name"] == "投标模板.docx"
    titles = [s["title"] for s in tpl["structure"]]
    assert any("技术方案要点" in t for t in titles)
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert detail["template"]["name"] == "投标模板.docx"


def test_bid_template_reject_non_docx(no_llm):
    # NO-009 FR-13: 非 docx 模板拒绝
    p = _mk_project("模板非docx")
    pid = p["id"]
    r = client.post(f"/api/bid/projects/{pid}/template",
                    files={"file": ("模板.pdf", b"%PDF", "application/pdf")})
    assert r.status_code == 400
    assert "docx" in r.json()["error"]


def test_bid_template_overwrite_and_delete(no_llm):
    # NO-009 FR-13: 单模板覆盖 + 删除
    from app.bidding.bid_engine import load_bid_template
    p = _mk_project("模板覆盖")
    pid = p["id"]
    tpl_path = _mk_template_docx()
    with open(tpl_path, "rb") as f:
        client.post(f"/api/bid/projects/{pid}/template",
                    files={"file": ("旧模板.docx", f.read(), "application/octet-stream")})
    with open(tpl_path, "rb") as f:
        r = client.post(f"/api/bid/projects/{pid}/template",
                        files={"file": ("新模板.docx", f.read(), "application/octet-stream")})
    assert r.json()["template"]["name"] == "新模板.docx"
    tpl = load_bid_template(pid)
    assert tpl["name"] == "新模板.docx"
    r = client.delete(f"/api/bid/projects/{pid}/template")
    assert r.json()["success"] is True
    assert load_bid_template(pid) is None


def test_bid_template_export_docx(no_llm):
    # NO-009 FR-14: 有模板导出 docx → 模板原内容保留 + 生成内容按章节插入
    import io
    import zipfile
    from app.bidding.bid_engine import save_bid_template
    p = _mk_project("模板化导出")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    doc = client.post(f"/api/bid/projects/{pid}/generate", json={"type": "tech_proposal"}).json()["doc"]
    tpl_path = _mk_template_docx()
    with open(tpl_path, "rb") as f:
        save_bid_template(pid, f.read(), "投标模板.docx")
    r = client.get(f"/api/bid/projects/{pid}/export/{doc['id']}?fmt=docx")
    assert r.status_code == 200
    with zipfile.ZipFile(io.BytesIO(r.content)) as z:
        docxml = z.read("word/document.xml").decode("utf-8", "ignore")
    assert "【模板保留】技术要点占位" in docxml, "模板原内容未保留"
    assert "【模板保留】项目理解占位" in docxml, "模板其他章节未保留"
    assert "CPU 主频" in docxml, "生成内容未插入模板对应章节"
    assert "ISO9001" in docxml, "生成内容未插入模板对应章节"


# ==================== FR-15 技术方案演示网页 ====================

def test_bid_demo_generate_html(no_llm):
    # NO-009 FR-15: tech_demo 生成自包含 html，含 KPI/图表/表格/告警与拆标报告驱动数据
    p = _mk_project("演示网页")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    r = client.post(f"/api/bid/projects/{pid}/generate", json={"type": "tech_demo"})
    assert r.status_code == 200, r.text
    doc = r.json()["doc"]
    assert doc["type"] == "tech_demo"
    assert doc["path"].endswith(".html")
    html = open(doc["path"], encoding="utf-8").read()
    for key in ("资源总览大屏", "class=\"kpi\"", "<svg", "<table", "最新告警",
                "技术参数响应清单", "模拟界面原型", "CPU 主频"):
        assert key in html, f"demo 缺少区块: {key}"


def test_bid_demo_export_html_and_shot_table(no_llm):
    # NO-009 FR-15: html 导出 + tech_proposal docx 功能截图区 Word 表格骨架
    import io
    import zipfile
    p = _mk_project("演示导出")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    demo = client.post(f"/api/bid/projects/{pid}/generate", json={"type": "tech_demo"}).json()["doc"]
    r = client.get(f"/api/bid/projects/{pid}/export/{demo['id']}?fmt=html")
    assert r.status_code == 200
    assert "text/html" in r.headers.get("content-type", "")
    tp = client.post(f"/api/bid/projects/{pid}/generate", json={"type": "tech_proposal"}).json()["doc"]
    r = client.get(f"/api/bid/projects/{pid}/export/{tp['id']}?fmt=docx")
    assert r.status_code == 200
    with zipfile.ZipFile(io.BytesIO(r.content)) as z:
        docxml = z.read("word/document.xml").decode("utf-8", "ignore")
    assert "功能截图区" in docxml
    assert "<w:tbl>" in docxml


# ==================== FR-16 项目资料入库 ====================

def test_bid_project_kb_write_on_upload(no_llm):
    # NO-009 FR-16: 上传规范书 → 原文件保留 + 文本写入项目级知识库，可检索命中
    from app import db as _db
    from app.bidding.bid_engine import ensure_project_kb, project_kb_name, get_project_kb_ids
    from app.knowledge import search_knowledge
    p = _mk_project("入库项目")
    pid = p["id"]
    _upload(pid, BID_TXT)
    kb_id = ensure_project_kb(pid)
    kbs = {kb["name"]: kb for kb in _db.db_list_knowledge_bases()}
    assert project_kb_name(pid) in kbs
    # 拆标后报告也入库
    client.post(f"/api/bid/projects/{pid}/parse")
    hits = search_knowledge("CPU 主频", [kb_id], top_k=3)
    assert hits, "项目知识库检索未命中"
    # 检索范围含项目库
    assert kb_id in get_project_kb_ids(pid)


# ==================== FR-17 需求分析 ====================

def test_bid_requirements_prd(no_llm):
    # NO-009 FR-17: 需求分析 → 结构化 PRD（roles/features/pages/interactions）+ 落库
    p = _mk_project("需求分析")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    r = client.post(f"/api/bid/projects/{pid}/requirements")
    assert r.status_code == 200, r.text
    prd = r.json()["prd"]
    for k in ("summary", "roles", "features", "pages", "interactions", "acceptance"):
        assert k in prd, f"PRD 缺少字段: {k}"
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert detail["prd_json"]["summary"]


def test_bid_requirements_without_parse_rejected(no_llm):
    # NO-009 FR-17: 未拆标时触发需求分析 → 400
    p = _mk_project("需求前置")
    pid = p["id"]
    r = client.post(f"/api/bid/projects/{pid}/requirements")
    assert r.status_code == 400


# ==================== FR-18 假页面生成 ====================

def test_bid_mockup_generate(no_llm):
    # NO-009 FR-18: 假页面 → 可预览 HTML 落盘 outputs/ + 成果列表 + 步骤状态
    p = _mk_project("假页面")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    client.post(f"/api/bid/projects/{pid}/requirements")
    r = client.post(f"/api/bid/projects/{pid}/mockup")
    assert r.status_code == 200, r.text
    doc = r.json()["doc"]
    assert doc["type"] == "tech_demo" and doc["kind"] == "mockup"
    assert os.path.isfile(doc["path"])
    html = open(doc["path"], encoding="utf-8").read()
    assert "模拟界面原型" in html
    assert "技术参数响应清单" in html
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert any(d["kind"] == "mockup" for d in detail["generated_docs"])


# ==================== FR-19 大纲与逐章生成 ====================

def test_bid_outline_and_chapter(no_llm):
    # NO-009 FR-19: 大纲生成 → 逐章生成（单章落库，source=rule 降级，失败不中断）
    p = _mk_project("大纲逐章")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    r = client.post(f"/api/bid/projects/{pid}/outline")
    assert r.status_code == 200, r.text
    outline = r.json()["outline"]
    assert len(outline) >= 8
    assert outline[0]["index"] == 1 and outline[0]["title"]
    r = client.post(f"/api/bid/projects/{pid}/chapter", json={"index": 1})
    assert r.status_code == 200, r.text
    ch = r.json()["chapter"]
    assert ch["index"] == 1 and ch["content"]
    assert ch["source"] in ("llm", "rule")
    # 未生成大纲 → 400
    p2 = _mk_project("无大纲")
    pid2 = p2["id"]
    r = client.post(f"/api/bid/projects/{pid2}/chapter", json={"index": 1})
    assert r.status_code == 400


# ==================== FR-20 左右对照确认 ====================

def test_bid_chapter_confirm(no_llm):
    # NO-009 FR-20: 确认本章 → 定稿锁定；重生成丢弃旧稿；仅确认章节参与组装
    p = _mk_project("章节确认")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    client.post(f"/api/bid/projects/{pid}/outline")
    client.post(f"/api/bid/projects/{pid}/chapter", json={"index": 1})
    # 确认（带编辑覆盖）
    r = client.post(f"/api/bid/projects/{pid}/chapters/confirm",
                    json={"index": 1, "content": "## 手动定稿内容"})
    assert r.status_code == 200, r.text
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert detail["chapters_json"][0]["confirmed"] is True
    assert "手动定稿内容" in detail["chapters_json"][0]["content"]
    # 无草稿章节不可确认
    r = client.post(f"/api/bid/projects/{pid}/chapters/confirm", json={"index": 2})
    assert r.status_code == 400


# ==================== FR-21 组装导出 ====================

def test_bid_assemble_document(no_llm):
    # NO-009 FR-21: 全部章节确认 → 组装 → 成果列表含组装文档，可导出 md
    p = _mk_project("组装项目")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    client.post(f"/api/bid/projects/{pid}/outline")
    outline = client.get(f"/api/bid/projects/{pid}").json()["project"]["outline_json"]
    for i in range(1, len(outline) + 1):
        client.post(f"/api/bid/projects/{pid}/chapter", json={"index": i})
        client.post(f"/api/bid/projects/{pid}/chapters/confirm", json={"index": i})
    r = client.post(f"/api/bid/projects/{pid}/assemble")
    assert r.status_code == 200, r.text
    doc = r.json()["doc"]
    assert doc["type"] == "tech_proposal" and doc["kind"] == "assemble"
    assert os.path.isfile(doc["path"])
    md = open(doc["path"], encoding="utf-8").read()
    assert "技术方案建议书" in md
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert any(d["kind"] == "assemble" for d in detail["generated_docs"])
    r = client.get(f"/api/bid/projects/{pid}/export/{doc['id']}?fmt=md")
    assert r.status_code == 200
    assert "text/markdown" in r.headers.get("content-type", "")


def test_bid_assemble_without_confirmed(no_llm):
    # NO-009 FR-21: 未确认任何章节 → 组装 400
    p = _mk_project("空组装")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    client.post(f"/api/bid/projects/{pid}/outline")
    r = client.post(f"/api/bid/projects/{pid}/assemble")
    assert r.status_code == 400


# ==================== FR-22/FR-23 一键智能起草流水线 ====================

def test_bid_pipeline_run(no_llm):
    # NO-009 FR-22: 一键起草（默认停人工复核）→ 拆标/需求/演示/大纲/逐章/截图全部产出
    p = _mk_project("一键起草")
    pid = p["id"]
    _upload(pid, BID_TXT)
    r = client.post(f"/api/bid/projects/{pid}/pipeline/run", json={})
    assert r.status_code == 200, r.text
    assert r.json()["status"] == "awaiting_review"
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert detail["parse_report"] and detail["parse_report"].get("qualifications")
    assert detail["prd_json"] and detail["prd_json"].get("features")
    assert any(d["kind"] == "mockup" for d in detail["generated_docs"])
    assert detail["outline_json"] and len(detail["outline_json"]) >= 8
    assert len(detail["chapters_json"]) == len(detail["outline_json"])
    # 进度状态：awaiting_review 且非 running
    st = client.get(f"/api/bid/projects/{pid}/pipeline/status").json()["status"]
    assert st["status"] == "awaiting_review"
    assert st["done"] is False and st["running"] is False


def test_bid_pipeline_auto_confirm(no_llm):
    # NO-009 FR-22: 全自动（auto_confirm）→ 自动确认全部章节 → 组装 → 自检 → 导出 docx
    p = _mk_project("全自动")
    pid = p["id"]
    _upload(pid, BID_TXT)
    r = client.post(f"/api/bid/projects/{pid}/pipeline/run", json={"auto_confirm": True})
    assert r.status_code == 200, r.text
    data = r.json()
    assert data["status"] == "completed"
    assert data["doc"]["kind"] == "assemble"
    assert data["docx"] and os.path.isfile(data["docx"])
    assert data["check"] and isinstance(data["check"]["unresponded"], list)
    detail = client.get(f"/api/bid/projects/{pid}").json()["project"]
    assert all(c["confirmed"] for c in detail["chapters_json"])
    assert detail["check_result"]
    st = client.get(f"/api/bid/projects/{pid}/pipeline/status").json()["status"]
    assert st["status"] == "completed" and st["done"] is True


def test_bid_pipeline_without_files_rejected(no_llm):
    # NO-009 FR-22: 未上传规范书 → 一键起草 400
    p = _mk_project("无文件流水线")
    pid = p["id"]
    r = client.post(f"/api/bid/projects/{pid}/pipeline/run", json={})
    assert r.status_code == 400
    assert "上传" in r.json()["error"]
    # 项目不存在 → 404
    r = client.post("/api/bid/projects/999999/pipeline/run", json={})
    assert r.status_code == 404


# ==================== FR-15 演示截图真插入 ====================

def _make_png(w=1, h=1):
    """构造最小有效 PNG（zlib 手写，免 PIL 依赖；python-docx add_picture 可识别）"""
    import struct
    import zlib

    def _chunk(typ, data):
        c = struct.pack(">I", len(data)) + typ + data
        return c + struct.pack(">I", zlib.crc32(typ + data) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", w, h, 8, 6, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\x00\x00\x00\x00" * (w * h))
    return (b"\x89PNG\r\n\x1a\n" + _chunk(b"IHDR", ihdr)
            + _chunk(b"IDAT", idat) + _chunk(b"IEND", b""))


def test_bid_screenshot_embed(no_llm, monkeypatch):
    # NO-009 FR-15: 有 playwright 截图清单 → 组装 docx 真插入图片（word/media/ 含图）；无截图不崩溃且无图
    import io
    import zipfile
    from app.bidding import screenshot as shot_mod
    from app.bidding.bid_engine import BID_UPLOAD_ROOT
    p = _mk_project("截图嵌入")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    client.post(f"/api/bid/projects/{pid}/outline")
    outline = client.get(f"/api/bid/projects/{pid}").json()["project"]["outline_json"]
    for i in range(1, len(outline) + 1):
        client.post(f"/api/bid/projects/{pid}/chapter", json={"index": i})
        client.post(f"/api/bid/projects/{pid}/chapters/confirm", json={"index": i})
    # 模拟已有截图：真实写入 outputs/mockup-shots/（局部 import 命中 patch 后的模块属性）
    shot_dir = os.path.join(BID_UPLOAD_ROOT, str(pid), "outputs", "mockup-shots")
    os.makedirs(shot_dir, exist_ok=True)
    png = os.path.join(shot_dir, "p1-full.png")
    with open(png, "wb") as f:
        f.write(_make_png())
    fake = [{"name": "全部页面", "file": "p1-full.png", "path": png, "section": "full"}]
    monkeypatch.setattr(shot_mod, "load_shots", lambda pid_: fake)
    doc = client.post(f"/api/bid/projects/{pid}/assemble").json()["doc"]
    r = client.get(f"/api/bid/projects/{pid}/export/{doc['id']}?fmt=docx")
    assert r.status_code == 200
    with zipfile.ZipFile(io.BytesIO(r.content)) as z:
        names = z.namelist()
        docxml = z.read("word/document.xml").decode("utf-8", "ignore")
    assert any(n.startswith("word/media/") and n.endswith(".png") for n in names), "docx 未嵌入截图"
    assert "blip" in docxml
    # 组装 md 亦含截图引用
    md = open(doc["path"], encoding="utf-8").read()
    assert "系统演示界面截图" in md and "p1-full.png" in md
    # 无截图 → 组装与导出不崩溃，且无嵌入图片
    monkeypatch.setattr(shot_mod, "load_shots", lambda pid_: [])
    doc2 = client.post(f"/api/bid/projects/{pid}/assemble").json()["doc"]
    r2 = client.get(f"/api/bid/projects/{pid}/export/{doc2['id']}?fmt=docx")
    assert r2.status_code == 200
    with zipfile.ZipFile(io.BytesIO(r2.content)) as z:
        names2 = z.namelist()
    assert not any(n.startswith("word/media/") and n.endswith(".png") for n in names2)
    # 无截图时导出 md 不含截图引用
    md2 = open(doc2["path"], encoding="utf-8").read()
    assert "系统演示界面截图" not in md2


def test_bid_shot_mockup_degrade_without_browser(no_llm):
    # NO-009 FR-15: playwright 缺失/渲染失败 → shot_mockup 降级不抛异常
    from app.bidding import screenshot as shot_mod
    import app.bidding.screenshot as _s
    p = _mk_project("截图降级")
    pid = p["id"]
    _upload(pid, BID_TXT)
    client.post(f"/api/bid/projects/{pid}/parse")
    client.post(f"/api/bid/projects/{pid}/mockup")
    # 无演示页 → 直接 degraded
    r = shot_mod.shot_mockup(999999)
    assert r["degraded"] is True and r["shots"] == []
    # 未安装 playwright → degraded 不抛异常
    import builtins
    real_import = builtins.__import__
    def fake_import(name, *a, **k):
        if name == "playwright.sync_api":
            raise ImportError("playwright 未安装")
        return real_import(name, *a, **k)
    monkeypatch = pytest.MonkeyPatch()
    monkeypatch.setattr(builtins, "__import__", fake_import)
    try:
        r = shot_mod.shot_mockup(pid)
    finally:
        monkeypatch.undo()
    assert r["degraded"] is True
    assert "playwright" in r["reason"]

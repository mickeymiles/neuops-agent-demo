# -*- coding: utf-8 -*-
"""本地 RAG 引擎：文档解析 / 切块 / Embedding / 向量索引 / 检索

依赖：
    chromadb      - 本地持久化向量库（chroma_data/ 目录）
    fastembed     - 本地中文 Embedding（BAAI/bge-small-zh-v1.5，onnxruntime 推理，无需 torch）
    openpyxl      - xlsx 解析
    pypdf         - pdf 解析
    python-docx   - docx 解析

容错设计：
    Chroma / fastembed 不可用时，search_knowledge 降级为
    SQLite knowledge_chunks 的关键词匹配（BM25 简化版），
    保证现有 mock 对话流程不被拖垮。
"""
import os
import re
import threading

# 模型下载镜像：HuggingFace 直连不稳定时自动走国内镜像（仅下载阶段生效）
os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")

from . import config

# Chroma 持久化目录
CHROMA_DIR = os.path.join(config.BASE_DIR, "chroma_data")
# 上传文件存放目录
UPLOAD_DIR = os.path.join(config.BASE_DIR, "uploads")

EMBED_MODEL = "BAAI/bge-small-zh-v1.5"
CHUNK_SIZE = 400      # 每块最大字符数
CHUNK_OVERLAP = 50    # 块间重叠字符数
DEFAULT_TOP_K = 5

_embed_lock = threading.Lock()
_embedder = None
_chroma_client = None


# ==================== 初始化 ====================

def _get_embedder():
    """单例获取 fastembed 向量模型；失败返回 None"""
    global _embedder
    if _embedder is not None:
        return _embedder
    with _embed_lock:
        if _embedder is not None:
            return _embedder
        try:
            from fastembed import TextEmbedding
            _embedder = TextEmbedding(model_name=EMBED_MODEL)
        except Exception as e:  # 模型下载失败 / 依赖缺失
            print(f"[knowledge] Embedding 初始化失败（将使用关键词降级检索）: {e}")
            _embedder = False
        return _embedder


def _embed_texts(texts):
    """批量向量化，返回 list[list[float]]；失败返回 None"""
    model = _get_embedder()
    if not model:
        return None
    try:
        return [list(map(float, v)) for v in model.embed(texts)]
    except Exception as e:
        print(f"[knowledge] Embedding 失败: {e}")
        return None


def _get_chroma():
    """单例获取 Chroma 客户端（PersistentClient）；失败返回 None"""
    global _chroma_client
    if _chroma_client is not None:
        return _chroma_client
    with _embed_lock:
        if _chroma_client is not None:
            return _chroma_client
        try:
            import chromadb
            os.makedirs(CHROMA_DIR, exist_ok=True)
            _chroma_client = chromadb.PersistentClient(path=CHROMA_DIR)
        except Exception as e:
            print(f"[knowledge] Chroma 初始化失败（将使用关键词降级检索）: {e}")
            _chroma_client = False
        return _chroma_client


def _get_collection():
    client = _get_chroma()
    if not client:
        return None
    try:
        return client.get_or_create_collection(
            name="knowledge_chunks",
            metadata={"hnsw:space": "cosine"},
        )
    except Exception as e:
        print(f"[knowledge] 获取集合失败: {e}")
        return None


# ==================== 文档解析 ====================

def parse_document(path):
    """按扩展名解析文档为纯文本；返回字符串（失败返回空串）"""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".txt" or ext == ".md":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif ext == ".xlsx":
            return _parse_xlsx(path)
        elif ext == ".pdf":
            return _parse_pdf(path)
        elif ext == ".docx":
            return _parse_docx(path)
        else:
            print(f"[knowledge] 不支持的扩展名: {ext}")
            return ""
    except Exception as e:
        print(f"[knowledge] 解析失败 {path}: {e}")
        return ""


def _parse_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        # 表头
        header = rows[0]
        lines.append(f"## 工作表: {ws.title}")
        for row in rows[1:]:
            cells = []
            for i, v in enumerate(row):
                if v is None:
                    continue
                col_name = header[i] if i < len(header) and header[i] else f"列{i+1}"
                cells.append(f"{col_name}: {v}")
            if cells:
                lines.append("；".join(cells))
    wb.close()
    return "\n".join(lines)


def _parse_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        t = page.extract_text() or ""
        if t.strip():
            pages.append(f"## 第{i+1}页\n{t}")
    return "\n".join(pages)


def _parse_docx(path):
    import docx
    d = docx.Document(path)
    parts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ==================== 切块 ====================

def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """按段落聚合 + 定长切块 + 重叠，返回 chunk 列表"""
    text = (text or "").strip()
    if not text:
        return []
    # 先按段落划分
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    # 将过短段落聚合成块
    groups = []
    buf = ""
    for p in paras:
        if len(buf) + len(p) > chunk_size and buf:
            groups.append(buf)
            buf = p
        else:
            buf = (buf + "\n" + p) if buf else p
    if buf:
        groups.append(buf)
    # 对每个组按定长二次切块（带重叠）
    chunks = []
    for g in groups:
        if len(g) <= chunk_size:
            chunks.append(g)
            continue
        start = 0
        while start < len(g):
            chunks.append(g[start:start + chunk_size])
            start += chunk_size - overlap
    # 过滤纯空白/过短碎片
    out = []
    for c in chunks:
        c = c.strip()
        if len(c) >= 8:
            out.append(c)
    return out


# ==================== 索引构建 ====================

def build_kb_index(kb_id: str, files):
    """将文件列表解析切块后写入向量库与 SQLite 元数据。

    files: list[path]（原文件路径，通常来自 uploads/{kb_id}/）
    返回 {"doc_count": n, "chunk_count": m, "vector": bool}
    """
    from . import db

    all_texts, all_metas = [], []
    doc_count = 0
    for path in files:
        text = parse_document(path)
        if not text.strip():
            continue
        doc_count += 1
        doc_name = os.path.basename(path)
        chunks = chunk_text(text)
        for idx, c in enumerate(chunks):
            all_texts.append(c)
            all_metas.append({"kb_id": kb_id, "doc_name": doc_name, "chunk_index": idx})

    # 清空旧索引（重建语义）
    db.db_clear_kb_chunks(kb_id)
    collection = _get_collection()
    if collection is not None:
        try:
            collection.delete(where={"kb_id": kb_id})
        except Exception:
            pass

    if not all_texts:
        db.db_update_kb_stats(kb_id, 0, 0)
        return {"doc_count": 0, "chunk_count": 0, "vector": False}

    # 写入 SQLite 元数据
    docs_by_name = {}
    for i, meta in enumerate(all_metas):
        docs_by_name.setdefault(meta["doc_name"], []).append(all_texts[i])
    for doc_name, chunks in docs_by_name.items():
        db.db_add_kb_chunks(kb_id, doc_name, chunks)

    # 写入向量库
    vectors = _embed_texts(all_texts)
    vector_ok = False
    if vectors is not None and collection is not None:
        try:
            ids = [f"{kb_id}-{i}" for i in range(len(all_texts))]
            collection.add(
                ids=ids,
                embeddings=vectors,
                documents=all_texts,
                metadatas=[{
                    "kb_id": kb_id,
                    "doc_name": m["doc_name"],
                    "chunk_index": m["chunk_index"],
                } for m in all_metas],
            )
            vector_ok = True
        except Exception as e:
            print(f"[knowledge] 写入向量库失败: {e}")

    db.db_update_kb_stats(kb_id, doc_count, len(all_texts))
    return {"doc_count": doc_count, "chunk_count": len(all_texts), "vector": vector_ok}


# ==================== 检索 ====================

def search_knowledge(query, kb_ids, top_k=DEFAULT_TOP_K):
    """检索知识库，返回 [{title, summary, score, source}]；score 为相似度 0~1

    优先向量检索；向量不可用时降级 SQLite 关键词匹配。
    """
    kb_ids = list(kb_ids or [])
    if not query or not kb_ids:
        return []

    # 1) 向量检索
    result = _vector_search(query, kb_ids, top_k)
    if result:
        return result
    # 2) 关键词降级
    return _keyword_search(query, kb_ids, top_k)


def _vector_search(query, kb_ids, top_k):
    collection = _get_collection()
    vectors = _embed_texts([query])
    if collection is None or not vectors:
        return []
    try:
        res = collection.query(
            query_embeddings=[vectors[0]],
            n_results=top_k * 2,
            where={"kb_id": {"$in": kb_ids}},
        )
    except Exception as e:
        print(f"[knowledge] 向量检索失败: {e}")
        return []

    out = []
    ids = res.get("ids") or [[]]
    docs = res.get("documents") or [[]]
    dists = res.get("distances") or [[]]
    metas = res.get("metadatas") or [[]]
    for i in range(len(ids[0])):
        score = max(0.0, min(1.0, 1.0 - float(dists[0][i])))  # cosine 距离转相似度
        meta = metas[0][i] or {}
        out.append({
            "title": meta.get("doc_name", ""),
            "summary": (docs[0][i] or "")[:200],
            "score": round(score, 4),
            "source": meta.get("doc_name", ""),
        })
    out.sort(key=lambda x: x["score"], reverse=True)
    return out[:top_k]


def _tokenize_zh(text):
    """中文分词（降级检索用）：标点切分 + 2-4 字滑动窗口，返回词频 dict"""
    parts = re.split(r"[\s，。、,.;；:：?!？！()（）\[\]【】\"'“”‘’<>《》/\\|_\-+*=~@#$%^&]+", text)
    freq = {}
    for p in parts:
        p = p.strip()
        if not p:
            continue
        if len(p) <= 4:
            freq[p] = freq.get(p, 0) + 2
        else:
            for n in (4, 3, 2):
                for i in range(len(p) - n + 1):
                    w = p[i:i + n]
                    if not re.search(r"[\u4e00-\u9fff]|[a-zA-Z0-9]", w):
                        continue
                    freq[w] = freq.get(w, 0) + n
    return freq


def _keyword_search(query, kb_ids, top_k):
    """SQLite 关键词匹配降级检索：按词频打分排序"""
    from . import db
    q_freq = _tokenize_zh(query)
    if not q_freq:
        return []
    q_words = sorted(q_freq, key=lambda w: -q_freq[w])[:20]

    scored = {}
    for kb_id in kb_ids:
        for r in db.db_list_kb_chunks(kb_id, limit=800):
            c = r["content"]
            score = 0.0
            hit = 0
            for w in q_words:
                if w in c:
                    score += q_freq[w]
                    hit += 1
            if hit:
                scored[r["id"]] = {
                    "title": r["doc_name"],
                    "summary": c[:200],
                    "score": round(min(0.99, 0.4 + score / 60.0), 4),
                    "source": r["doc_name"],
                    "_hits": hit,
                }
    ordered = sorted(scored.values(), key=lambda x: (-x["_hits"], -x["score"]))
    for o in ordered:
        o.pop("_hits", None)
    return ordered[:top_k]


# ==================== 知识库级操作 ====================

def delete_kb_vectors(kb_id: str):
    """删除某知识库的全部向量（删除知识库时调用）"""
    collection = _get_collection()
    if collection is None:
        return
    try:
        collection.delete(where={"kb_id": kb_id})
    except Exception as e:
        print(f"[knowledge] 删除向量失败: {e}")


def delete_chunk_vector(kb_id: str, chunk_index: int, doc_name: str):
    """删除单条 chunk 向量（尽力而为；找不到就忽略）"""
    collection = _get_collection()
    if collection is None:
        return
    try:
        collection.delete(where={
            "$and": [
                {"kb_id": kb_id},
                {"doc_name": doc_name},
                {"chunk_index": chunk_index},
            ]
        })
    except Exception:
        pass


def rebuild_stats(kb_id: str):
    """重建元数据统计（doc_count / chunk_count）"""
    from . import db
    chunks = db.db_list_kb_chunks(kb_id, limit=100000)
    doc_names = {c["doc_name"] for c in chunks}
    db.db_update_kb_stats(kb_id, len(doc_names), len(chunks))


def get_upload_dir(kb_id: str):
    d = os.path.join(UPLOAD_DIR, kb_id)
    os.makedirs(d, exist_ok=True)
    return d

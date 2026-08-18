# -*- coding: utf-8 -*-
"""投标业务引擎：拆标(规则粗筛+LLM精提炼) / 生成(模板+知识库RAG) / 自检(清单核对) / 导出

复用现有体系：
- LLM：DeepSeek（读取与 agent_chat 相同 key，调用轨迹写入 llm_calls 观测表）
- RAG：app.knowledge.search_knowledge（Chroma 向量检索，降级关键词）
- 文本抽取：app.knowledge.parse_document（docx/pdf/xlsx/txt）
"""
import json
import os
import re
import time
import uuid

import httpx

from .. import config, db, knowledge

# 投标工作台成果根目录
BID_UPLOAD_ROOT = os.path.join(config.BASE_DIR, "uploads", "bid")

# 投标知识库名称（预置 5 类，seed 时按 name 建立；id 由 db_create_knowledge_base 生成）
BID_KB_NAMES = ("投标-资质库", "投标-业绩库", "投标-素材库", "投标-模板库", "投标-人员库")

# 拆标六类章节 → 正则关键词（规则粗筛用）
SECTION_KEYWORDS = {
    "qualifications": ["资格要求", "资质要求", "资格条件", "准入条件", "投标人资格", "投标人资质", "资格标准"],
    "performance": ["业绩要求", "类似项目", "同类项目", "业绩证明", "成功案例", "项目案例"],
    "tech_params": ["技术参数", "技术需求", "货物需求", "功能需求", "技术指标", "★", "▲", "关键技术"],
    "scoring": ["评分标准", "评分细则", "评审标准", "评标办法", "评标细则", "评分办法", "得分"],
    "rejection_clauses": ["废标", "否决", "无效投标", "作废", "拒绝投标", "取消资格"],
}

LLM_CHUNK_SIZE = 6000  # 拆标分块（字符）

# 招标文件分类（上传打标签，FR-10）
BID_FILE_CATEGORIES = {
    "feasibility": "可研文件",
    "technical": "技术要求",
    "commercial": "商务要求",
    "contract": "合同文稿",
    "format": "格式要求",
    "other": "其他",
}
# 拆标主文本类别（优先级从高到低）；可研单独隔离，避免过时参数污染正式标书
MAIN_PARSE_CATEGORIES = ("technical", "commercial", "contract", "other")

# 自动识别：文件名关键词（精确）→ 内容兜底（保守，见 auto_category）
_FILE_NAME_RULES = (
    (("可研", "可行性研究", "项目建议书", "立项报告", "初步设计"), "feasibility"),
    (("格式要求", "编制要求", "装订要求", "密封要求", "投标文件格式", "文档格式", "排版要求", "格式模板"), "format"),
    (("技术规范", "技术需求", "技术要求", "技术参数", "货物需求", "服务需求", "采购需求"), "technical"),
    (("商务要求", "商务条款", "投标报价", "报价表"), "commercial"),
    (("合同", "协议书", "合同文本", "合同草案"), "contract"),
)
_CONTENT_RULES = (
    (r"格式要求|编制要求|装订要求|密封要求|页边距|行距|字体|字号|排版", "format"),
    (r"合同条款|合同正文|协议书", "contract"),
    (r"技术规范|技术需求|技术参数|技术指标|关键参数|★|▲", "technical"),
    (r"投标报价|商务条款|报价表", "commercial"),
    (r"可行性研究|项目建议书|立项", "feasibility"),
)


def auto_category(name: str, head_text: str = "") -> str:
    """自动识别文件分类（仅作初值，用户可在界面手动覆盖；FR-10）"""
    for kws, cat in _FILE_NAME_RULES:
        if any(k in (name or "") for k in kws):
            return cat
    head = (head_text or "")[:500]
    for pat, cat in _CONTENT_RULES:
        if re.search(pat, head):
            return cat
    return "other"

# 生成材料类型
DOC_TYPES = {
    "tech_proposal": "技术方案建议书",
    "response": "招标点对点应答",
    "ppt_outline": "售前汇报PPT大纲",
    "impl_plan": "运维实施方案",
    "tech_demo": "技术方案演示网页",
}

LLM_JSON_SCHEMA_HINT = """输出严格 JSON（不要输出其他内容），格式：
{
  "qualifications": [{"item": "资质名称", "level": "等级要求", "required_docs": "所需证明材料", "source": "原文出处"}],
  "performance": [{"item": "业绩要求描述", "requirement": "量化要求", "evidence": "证明材料", "source": "原文出处"}],
  "tech_params": [{"item": "参数项", "value": "要求值", "acceptance": "验收方式", "is_key": true/false, "source": "原文出处"}],
  "scoring": [{"item": "评分项", "score": 分值, "notes": "评分说明", "source": "原文出处"}],
  "rejection_clauses": ["废标/否决条款原文摘要"],
  "response_checklist": [{"id": "R1", "item": "应答材料项", "doc_type": "材料类型", "status": "todo", "source": "原文出处"}]
}
source 字段必须标注原文出处，格式「文件名·章节」，例如「技术规范书.docx·第3章 技术要求」；
若对应文本片段带有【来源文件：xxx】标记，请以该标记中的文件名为准。
"""


# ────────────────────────────────────────────
# 知识库 id 解析
# ────────────────────────────────────────────

def get_bid_kb_ids():
    """按名称解析 5 类投标知识库 id（未创建则返回已存在的子集）"""
    ids = []
    for name in BID_KB_NAMES:
        for kb in db.db_list_knowledge_bases():
            if kb.get("name") == name:
                ids.append(kb["id"])
                break
    return ids


# ────────────────────────────────────────────
# 项目级知识库（NO-009 FR-16）
# ────────────────────────────────────────────

def project_kb_name(project_id) -> str:
    """项目级知识库名称（唯一，用于幂等查找）"""
    return f"投标-项目{project_id}资料库"


def ensure_project_kb(project_id) -> str:
    """幂等获取/创建项目级知识库 bid-project-{pid}，返回 kb_id"""
    name = project_kb_name(project_id)
    for kb in db.db_list_knowledge_bases():
        if kb.get("name") == name:
            return kb["id"]
    kb_id = db.db_create_knowledge_base(
        name, f"项目 {project_id} 上传的规范书 / 拆标报告 / 模板（自动入库）")
    return kb_id


def get_project_kb_ids(project_id):
    """检索范围 = 项目级知识库 ∪ 预置 5 类投标库（FR-16）"""
    ids = get_bid_kb_ids()
    kb_id = ensure_project_kb(project_id)
    if kb_id and kb_id not in ids:
        ids.append(kb_id)
    return ids


def _kb_docs_dir(project_id) -> str:
    """项目知识库源文档目录（每份文本一个 .txt，全量重建）"""
    d = os.path.join(BID_UPLOAD_ROOT, str(project_id), "kb_docs")
    os.makedirs(d, exist_ok=True)
    return d


def _kb_write_text(project_id, doc_name: str, text: str) -> int:
    """把一段文本写入项目知识库（全量重建：重灌 kb_docs/ 下全部文档）"""
    if not (text or "").strip():
        return 0
    safe = _safe_name(doc_name) or "doc.txt"
    if not safe.endswith(".txt"):
        safe += ".txt"
    d = _kb_docs_dir(project_id)
    with open(os.path.join(d, safe), "w", encoding="utf-8") as f:
        f.write(text)
    files = [os.path.join(d, fn) for fn in sorted(os.listdir(d))
             if os.path.isfile(os.path.join(d, fn))]
    kb_id = ensure_project_kb(project_id)
    knowledge.build_kb_index(kb_id, files)
    return len(files)


def kb_write_upload_texts(project_id) -> None:
    """上传/拆标/模板后统一重灌：抽取文本 + 拆标报告 + 模板章节（FR-16）"""
    # 1) 规范书抽取文本（technical/commercial/contract/format/other）
    text = _read_extracted_text(project_id)
    if text.strip():
        _kb_write_text(project_id, "规范书抽取文本.md", text)
    # 2) 拆标报告
    proj = db.bid_get_project(project_id)
    if proj and (proj.get("parse_report") or {}):
        _kb_write_text(project_id, "拆标报告.md", _report_to_md(proj["parse_report"]))
    # 3) 模板章节结构
    tpl = load_bid_template(project_id)
    if tpl and tpl.get("structure"):
        lines = ["# 投标模板章节结构"]
        for s in tpl["structure"]:
            lines.append(f"{'#' * s['level']} {s['title']}")
        _kb_write_text(project_id, "投标模板章节.md", "\n".join(lines))


def _report_to_md(report: dict) -> str:
    """拆标报告六类 → 结构化 Markdown（供入库与逐章输入）"""
    if not report:
        return ""
    lines = ["# 拆标报告", ""]
    sec_titles = {
        "qualifications": "资质要求",
        "performance": "业绩要求",
        "tech_params": "技术参数",
        "scoring": "评分标准",
        "rejection_clauses": "废标/否决条款",
        "response_checklist": "应答清单",
    }
    for key, title in sec_titles.items():
        items = report.get(key) or []
        if not items:
            continue
        lines.append(f"## {title}")
        if key == "rejection_clauses":
            for c in items:
                lines.append(f"- {c}")
        else:
            for it in items:
                if isinstance(it, dict):
                    kv = "；".join(f"{k}: {v}" for k, v in it.items() if v and str(v) != "未识别")
                    lines.append(f"- {kv}")
                else:
                    lines.append(f"- {it}")
        lines.append("")
    return "\n".join(lines)


# ────────────────────────────────────────────
# 规则粗筛（不依赖 LLM，TC-2 兜底）
# ────────────────────────────────────────────

def _section_candidates(text: str) -> dict:
    """按行关键词匹配，产出六类骨架：每类最多 N 条候选句"""
    out = {k: [] for k in SECTION_KEYWORDS}
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    for k, kws in SECTION_KEYWORDS.items():
        hits = [ln[:120] for ln in lines if any(kw in ln for kw in kws)]
        # 去重保序
        seen, kept = set(), []
        for h in hits:
            if h not in seen:
                seen.add(h)
                kept.append(h)
            if len(kept) >= 6:
                break
        out[k] = kept
    return out


def rule_parse(text: str) -> dict:
    """规则粗筛：产出六类结构化骨架，缺项标「未识别」"""
    cand = _section_candidates(text)
    quals = [{"item": c, "level": "未识别", "required_docs": "未识别", "source": "规则粗筛"}
             for c in cand["qualifications"]] or \
            [{"item": "未识别", "level": "未识别", "required_docs": "未识别", "source": "规则粗筛"}]
    perfs = [{"item": c, "requirement": "未识别", "evidence": "未识别", "source": "规则粗筛"}
             for c in cand["performance"]] or \
            [{"item": "未识别", "requirement": "未识别", "evidence": "未识别", "source": "规则粗筛"}]
    techs = [{"item": c, "value": "未识别", "acceptance": "未识别",
              "is_key": ("★" in c or "▲" in c), "source": "规则粗筛"}
             for c in cand["tech_params"]] or \
            [{"item": "未识别", "value": "未识别", "acceptance": "未识别",
              "is_key": False, "source": "规则粗筛"}]
    scorings = [{"item": c, "score": 0, "notes": "未识别", "source": "规则粗筛"}
                for c in cand["scoring"]] or \
               [{"item": "未识别", "score": 0, "notes": "未识别", "source": "规则粗筛"}]
    checklist = [{"id": f"R{i+1}", "item": c, "doc_type": "证明文件",
                  "status": "todo", "source": "规则粗筛"}
                 for i, c in enumerate(cand["qualifications"][:3] + cand["performance"][:2])] or \
                [{"id": "R1", "item": "未识别", "doc_type": "未识别",
                  "status": "todo", "source": "规则粗筛"}]
    return {
        "qualifications": quals,
        "performance": perfs,
        "tech_params": techs,
        "scoring": scorings,
        "rejection_clauses": cand["rejection_clauses"] or ["未识别"],
        "response_checklist": checklist,
    }


# ────────────────────────────────────────────
# LLM 精提炼
# ────────────────────────────────────────────

def _load_deepseek_key():
    from ..agent_chat import _load_deepseek_key as _k
    return _k()


def _call_llm_json(system_prompt: str, user_text: str, trace_ctx=None):
    """调用 DeepSeek 返回 JSON dict；失败返回 None（调用轨迹写入观测表）"""
    key = _load_deepseek_key()
    if not key:
        return None
    payload = {
        "model": config.DEEPSEEK_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_text[:LLM_CHUNK_SIZE]},
        ],
        "temperature": 0.2,
        "max_tokens": 4000,
    }
    _t0 = time.time()
    try:
        r = httpx.post(
            f"{config.DEEPSEEK_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json=payload,
            timeout=180,
        )
        if r.status_code != 200:
            _record_llm(trace_ctx, payload, None, time.time() - _t0, error=f"HTTP {r.status_code}")
            return None
        data = r.json()
        _record_llm(trace_ctx, payload, data, time.time() - _t0)
        content = data["choices"][0]["message"]["content"]
        content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content.strip(), flags=re.S)
        return json.loads(content)
    except Exception as e:
        _record_llm(trace_ctx, payload, None, time.time() - _t0, error=str(e))
        return None


def _call_llm_text(system_prompt: str, user_text: str, trace_ctx=None):
    """调用 DeepSeek 返回纯文本 Markdown；失败返回空串"""
    key = _load_deepseek_key()
    if not key:
        return ""
    payload = {
        "model": config.DEEPSEEK_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_text[:LLM_CHUNK_SIZE * 2]},
        ],
        "temperature": 0.4,
        "max_tokens": 4000,
    }
    _t0 = time.time()
    try:
        r = httpx.post(
            f"{config.DEEPSEEK_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json=payload,
            timeout=180,
        )
        if r.status_code != 200:
            _record_llm(trace_ctx, payload, None, time.time() - _t0, error=f"HTTP {r.status_code}")
            return ""
        data = r.json()
        _record_llm(trace_ctx, payload, data, time.time() - _t0)
        content = data["choices"][0]["message"]["content"] or ""
        content = re.sub(r"^```(?:markdown|md)?\s*|\s*```$", "", content.strip(), flags=re.S)
        return content.strip()
    except Exception as e:
        _record_llm(trace_ctx, payload, None, time.time() - _t0, error=str(e))
        return ""


def _record_llm(trace_ctx, payload, resp, latency_s, error=""):
    """写 llm_calls 观测表（复用 agent_chat 的记录函数）"""
    try:
        from ..agent_chat import _record_llm_call as _rec
        _rec(trace_ctx, payload.get("model"), payload, resp, latency_s, error=error)
    except Exception:
        pass


def llm_refine(text: str):
    """分块喂 DeepSeek 按 JSON Schema 提炼；解析失败返回 None"""
    """分块喂 DeepSeek 按 JSON Schema 提炼；解析失败返回 None"""
    chunks = [text[i:i + LLM_CHUNK_SIZE] for i in range(0, len(text or ""), LLM_CHUNK_SIZE)] or [""]
    merged = None
    for i, chunk in enumerate(chunks):
        sys_p = "你是资深招投标解析专家。请从招标文件中提取结构化拆标信息。" + LLM_JSON_SCHEMA_HINT
        got = _call_llm_json(sys_p, chunk, trace_ctx={"stage": "bid_parse", "round_no": i})
        if not got:
            continue
        if merged is None:
            merged = got
        else:
            for k in merged:
                if isinstance(merged[k], list) and isinstance(got.get(k), list):
                    merged[k].extend(got[k])
    return merged


# ────────────────────────────────────────────
# 拆标入口
# ────────────────────────────────────────────

def parse_bid_document(project_id: int) -> dict:
    """拆标：主拆标仅合并正式标书类别（可研隔离），格式要求单独提取排版规范（FR-10/FR-12）"""
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    text = _read_extracted_text(project_id, categories=MAIN_PARSE_CATEGORIES)
    if not text.strip():
        # 兜底：主文本为空（如仅可研文件）→ 全量合并，保证可拆标
        text = _read_extracted_text(project_id)
    # 格式要求单独提取 → 排版规范（FR-12），docx 导出时覆盖默认基线
    fmt_text = _read_extracted_text(project_id, categories=("format",))
    if fmt_text.strip():
        save_format_spec(project_id, extract_format_spec(fmt_text))
    rule = rule_parse(text or "")
    refined = llm_refine(text or "")
    if refined:
        # LLM 结果优先，空章节回退规则骨架
        for k in rule:
            if k in refined and refined[k]:
                rule[k] = refined[k]
    # 兜底：保证六类齐全
    for k in rule:
        if not rule[k]:
            rule[k] = [{"item": "未识别", "source": "未识别"}]
    db.bid_save_parse_report(project_id, rule)
    db.bid_set_status(project_id, "已拆标")
    # 拆标报告同步写入项目知识库（FR-16）
    try:
        kb_write_upload_texts(project_id)
    except Exception:
        pass
    rule["format_spec"] = load_format_spec(project_id) or dict(DOCX_STYLE_DEFAULTS)
    return rule


def _read_extracted_text(project_id, categories=None) -> str:
    """读取抽取文本（按分类过滤），合并时带【来源文件】头便于溯源。
    - 新结构：extracted/<category>/<name>.txt
    - 旧结构：extracted/<name>.txt（平铺，视为 other，兼容旧数据）
    """
    ext_dir = os.path.join(BID_UPLOAD_ROOT, str(project_id), "extracted")
    parts = []
    if not os.path.isdir(ext_dir):
        return ""
    for entry in sorted(os.listdir(ext_dir)):
        p = os.path.join(ext_dir, entry)
        if os.path.isdir(p):
            cat = entry if entry in BID_FILE_CATEGORIES else "other"
            if categories is not None and cat not in categories:
                continue
            for fn in sorted(os.listdir(p)):
                body = _safe_read(os.path.join(p, fn))
                if body:
                    parts.append(f"【来源文件：{_pretty_src(fn)}】\n{body}")
        else:
            if categories is not None and "other" not in categories:
                continue
            body = _safe_read(p)
            if body:
                parts.append(f"【来源文件：{_pretty_src(entry)}】\n{body}")
    return "\n\n".join(parts)


def _safe_read(path: str) -> str:
    try:
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception:
        return ""


def _pretty_src(fn: str) -> str:
    """抽取文本名 → 可读来源名（去 .txt 后缀）"""
    return os.path.splitext(os.path.basename(fn))[0]


# ────────────────────────────────────────────
# 生成材料（模板 + 拆标报告 + 知识库 RAG）
# ────────────────────────────────────────────

def generate_document(project_id: int, doc_type: str) -> dict:
    """生成材料，落盘 outputs/{doc_id}.md，返回 doc 记录 dict"""
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    if doc_type not in DOC_TYPES:
        raise ValueError(f"不支持的材料类型: {doc_type}")
    report = proj.get("parse_report") or {}
    kb_ids = get_project_kb_ids(project_id)  # 项目库 ∪ 预置 5 类（FR-16）
    rag_lines = []
    for q in ("资质要求", "业绩要求", "技术方案要点", "实施方案", "人员配置"):
        for hit in knowledge.search_knowledge(q, kb_ids, top_k=2):
            rag_lines.append(f"- [{hit.get('title')}] {hit.get('summary')}")
    rag_text = "\n".join(rag_lines) or "（知识库暂无相关内容）"

    doc_id = uuid.uuid4().hex[:8]
    title = DOC_TYPES[doc_type]
    if doc_type == "tech_demo":
        # 技术方案演示网页（FR-15）：由拆标报告技术参数驱动生成自包含 HTML
        technical_text = _read_extracted_text(project_id, categories=("technical",))
        features = _demo_features(report, technical_text)
        body = _build_demo_html(proj, features, report.get("tech_params") or [])
        ext = "html"
    else:
        body = _build_doc_body(doc_type, proj, report, rag_text)
        ext = "md"

    output_dir = os.path.join(BID_UPLOAD_ROOT, str(project_id), "outputs")
    os.makedirs(output_dir, exist_ok=True)
    md_path = os.path.join(output_dir, f"{doc_id}.{ext}")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(body)

    doc = {
        "id": doc_id,
        "type": doc_type,
        "title": title,
        "path": md_path,
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    db.bid_add_generated_doc(project_id, doc)
    db.bid_set_status(project_id, "已生成")
    return doc


def _build_doc_body(doc_type, proj, report, rag_text) -> str:
    """按类型组装文档正文（无依据项输出【待补充材料】占位，不编造参数）"""
    name = proj.get("name", "本项目")
    tenderee = proj.get("tenderee", "招标方")
    lines = [f"# {name} — {DOC_TYPES[doc_type]}", f"招标方：{tenderee}", ""]

    if doc_type == "tech_proposal":
        lines += ["## 一、项目理解与总体方案", "【待补充材料】请依据拆标报告技术参数逐项展开。", ""]
        lines += ["## 二、技术方案要点"]
        for t in report.get("tech_params") or []:
            item, val = t.get("item", "未识别"), t.get("value", "未识别")
            lines.append(f"- {item}：{val}（验收：{t.get('acceptance', '待补充')}）")
        lines += ["", "## 三、实施方案概述", "【待补充材料】实施周期、资源投入、风险控制待补充。", ""]

    elif doc_type == "response":
        lines += ["## 一、资质响应"]
        for q in report.get("qualifications") or []:
            lines.append(f"- 资质：{q.get('item', '未识别')}（等级 {q.get('level', '未识别')}）→ 我方响应：{q.get('required_docs', '待补充材料')}")
        lines += ["", "## 二、业绩响应"]
        for p in report.get("performance") or []:
            lines.append(f"- {p.get('item', '未识别')}（{p.get('requirement', '未识别')}）→ 我方业绩：{p.get('evidence', '待补充材料')}")
        lines += ["", "## 三、技术参数点对点应答"]
        for t in report.get("tech_params") or []:
            lines.append(f"- {t.get('item', '未识别')}：要求 {t.get('value', '未识别')} → 我方应答：【待补充材料】")

    elif doc_type == "ppt_outline":
        lines += ["## 封面", "项目名称 / 投标单位 / 日期", "",
                  "## 目录", "1. 公司简介  2. 项目理解  3. 技术方案  4. 实施计划  5. 团队与资质  6. 服务承诺", "",
                  "## 1 公司简介", "【待补充材料】公司规模、行业资质、核心优势。", "",
                  "## 2 项目理解", f"面向 {tenderee} 的 {name}，核心诉求见拆标报告评分标准。", "",
                  "## 3 技术方案", "依据技术参数逐项阐述（见拆标报告）。", "",
                  "## 4 实施计划", "【待补充材料】里程碑、资源、风险预案。", "",
                  "## 5 团队与资质", "【待补充材料】项目团队与资质证书清单。", "",
                  "## 6 服务承诺", "【待补充材料】SLA、驻场、售后承诺。"]

    elif doc_type == "impl_plan":
        lines += ["## 一、实施目标", f"围绕 {name} 建设目标，提供从部署到运维的全周期服务。", "",
                  "## 二、实施内容", "【待补充材料】部署架构、环境要求、集成范围。", "",
                  "## 三、项目里程碑", "【待补充材料】分阶段里程碑与交付物。", "",
                  "## 四、运维服务方案", "【待补充材料】巡检、告警、应急、变更管理流程。", "",
                  "## 五、人员配置", "【待补充材料】项目经理、实施工程师、驻场运维人员。", ""]

    lines += ["", "---", "", "## 知识库参考", rag_text]
    lines += ["", "> 本文档由投标工作台生成，为 AI 初稿，请人工复核后使用。"]
    return "\n".join(lines)


# ────────────────────────────────────────────
# 合规自检（清单核对）
# ────────────────────────────────────────────

def check_compliance(project_id: int) -> dict:
    """自检：未响应项 / 废标红线 / 评分点得分建议"""
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    report = proj.get("parse_report") or {}
    docs = proj.get("generated_docs") or []
    text = ""
    for d in docs:
        p = d.get("path", "")
        if os.path.isfile(p):
            try:
                with open(p, encoding="utf-8", errors="ignore") as f:
                    text += f.read() + "\n"
            except Exception:
                continue

    result = {"unresponded": [], "redlines": [], "scoring": []}

    # 1) 未响应项：应答清单条目是否在生成文本中覆盖
    for c in report.get("response_checklist") or []:
        item = c.get("item", "")
        if not item or item == "未识别":
            continue
        keyword = re.sub(r"[【】（）()\s]", "", item)[:12]
        if keyword and keyword not in re.sub(r"[【】（）()\s]", "", text):
            result["unresponded"].append({"id": c.get("id"), "item": item, "reason": "生成材料中未找到对应内容"})

    # 2) 废标红线：逐条检查是否有应对说明
    for clause in report.get("rejection_clauses") or []:
        if clause == "未识别":
            continue
        kw = re.sub(r"[【】（）()\s]", "", clause)[:10]
        if kw and kw not in re.sub(r"[【】（）()\s]", "", text):
            result["redlines"].append({"clause": clause, "level": "高危", "advice": "必须在应答中正面应对，否则有废标风险"})

    # 3) 评分点得分建议：按覆盖度给分
    for s in report.get("scoring") or []:
        item = s.get("item", "")
        full = s.get("score") or 0
        if not item or item == "未识别":
            continue
        kw = re.sub(r"[【】（）()\s]", "", item)[:12]
        covered = bool(kw) and kw in re.sub(r"[【】（）()\s]", "", text)
        result["scoring"].append({
            "item": item,
            "full_score": full,
            "suggest_score": round(full * 0.85) if covered else 0,
            "advice": "已覆盖，建议补充量化证据" if covered else "未覆盖，需补充该评分点内容",
        })

    db.bid_save_check_result(project_id, result)
    db.bid_set_status(project_id, "已自检")
    return result


# ────────────────────────────────────────────
# 投标格式规范（FR-11/FR-12）
# ────────────────────────────────────────────

# 默认排版基线：正文宋体小四 1.5 倍行距、黑体标题、规范页边距、页脚页码、封面与目录
DOCX_STYLE_DEFAULTS = {
    "font_body": "宋体",
    "font_heading": "黑体",
    "size_body": 12,           # 小四
    "size_h1": 16,             # 三号
    "size_h2": 14,             # 四号
    "size_h3": 12,             # 小四
    "line_spacing": 1.5,
    "first_line_indent_cm": 0.85,  # 首行缩进 2 字符（12pt）
    "margin_top_cm": 2.54,
    "margin_bottom_cm": 2.54,
    "margin_left_cm": 3.17,
    "margin_right_cm": 3.17,
    "cover": True,
    "toc": True,
    "page_number": True,
}

FONT_SIZE_MAP = {"初号": 42, "小初": 36, "一号": 26, "小一": 24, "二号": 22, "小二": 18,
                 "三号": 16, "小三": 15, "四号": 14, "小四": 12, "五号": 10.5, "小五": 9}


def extract_format_spec(text: str) -> dict:
    """从格式要求文本关键词提取排版规范（FR-12）；未命中项不输出，由默认基线兜底"""
    spec = {}
    t = text or ""
    fonts = re.findall(r"(宋体|仿宋_?GB2312|仿宋|黑体|楷体|楷体_?GB2312|微软雅黑|方正小标宋|方正仿宋)", t)
    if fonts:
        spec["fonts"] = list(dict.fromkeys(fonts))
    m = re.search(r"(小初|小一|小二|小三|小四|小五|初号|一号|二号|三号|四号|五号)", t)
    if m:
        spec["size_body"] = FONT_SIZE_MAP[m.group(1)]
    m = re.search(r"(\d+(?:\.\d+)?)\s*倍\s*行距", t)
    if m:
        spec["line_spacing"] = float(m.group(1))
    m = re.search(r"行距\s*固定值\s*(\d+(?:\.\d+)?)\s*磅", t)
    if m:
        spec["line_spacing_pt"] = float(m.group(1))
    m = re.search(r"页边距\s*[：:]\s*上\s*([\d.]+)\s*(?:cm|毫米|mm).{0,8}下\s*([\d.]+)\s*(?:cm|毫米|mm)"
                  r".{0,8}左\s*([\d.]+)\s*(?:cm|毫米|mm).{0,8}右\s*([\d.]+)\s*(?:cm|毫米|mm)", t)
    if m:
        spec["margin_top_cm"], spec["margin_bottom_cm"] = float(m.group(1)), float(m.group(2))
        spec["margin_left_cm"], spec["margin_right_cm"] = float(m.group(3)), float(m.group(4))
    if re.search(r"页码|页脚", t):
        spec["page_number"] = True
    return spec


def save_format_spec(project_id, spec: dict) -> None:
    """写 format_spec.json（拆标时生成，供导出渲染）"""
    try:
        path = os.path.join(BID_UPLOAD_ROOT, str(project_id), "format_spec.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(spec, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def load_format_spec(project_id) -> dict:
    """读 format_spec.json；缺失返回空 dict（用默认基线）"""
    try:
        with open(os.path.join(BID_UPLOAD_ROOT, str(project_id), "format_spec.json"),
                  encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def _merge_format_spec(extracted: dict) -> dict:
    """格式要求覆盖默认基线（FR-12）：命中项优先，未命中用默认"""
    spec = dict(DOCX_STYLE_DEFAULTS)
    if not extracted:
        return spec
    if extracted.get("fonts"):
        spec["font_body"] = extracted["fonts"][0]
    if extracted.get("size_body"):
        spec["size_body"] = extracted["size_body"]
    if extracted.get("line_spacing"):
        spec["line_spacing"] = extracted["line_spacing"]
    for k in ("margin_top_cm", "margin_bottom_cm", "margin_left_cm", "margin_right_cm"):
        if extracted.get(k):
            spec[k] = extracted[k]
    if "page_number" in extracted:
        spec["page_number"] = bool(extracted["page_number"])
    return spec


# ────────────────────────────────────────────
# 投标模板（FR-13/FR-14）
# ────────────────────────────────────────────

_CN_SEQ_RE = re.compile(r"^[（(]?[一二三四五六七八九十百]+[）)、.．]\s*|^第[一二三四五六七八九十百]+[章节部分篇]")


def _safe_name(name: str) -> str:
    """文件名安全化：仅保留中文/字母/数字/._-，防路径穿越"""
    base = os.path.basename((name or "").replace("\\", "/")).strip()
    base = re.sub(r"[^\w\u4e00-\u9fff.\-]", "_", base)
    return base or "template.docx"


def _norm_title(s: str) -> str:
    """标题归一化：去'第X章/一、二、三、/'序号前缀与空白、小写，用于双向包含匹配"""
    t = (s or "").strip()
    t = re.sub(r"^第[一二三四五六七八九十百]+[章节部分篇]", "", t)
    t = re.sub(r"^[（(]?[一二三四五六七八九十百]+[）)、.．]\s*", "", t)
    t = re.sub(r"^\d+[\.、．)\s]+", "", t)
    t = re.sub(r"\s+", "", t)
    return t.lower()


def _template_dir(project_id) -> str:
    return os.path.join(BID_UPLOAD_ROOT, str(project_id), "template")


def parse_template_structure(template_path: str) -> list:
    """解析 docx 章节树：Heading 1/2/3 与带中文序号的段落（FR-13）"""
    out = []
    try:
        from docx import Document
        doc = Document(template_path)
    except Exception:
        return out
    seen = set()
    for p in doc.paragraphs:
        style = (p.style.name or "").lower() if p.style else ""
        text = (p.text or "").strip()
        if not text:
            continue
        level = None
        if style.startswith("heading 1"):
            level = 1
        elif style.startswith("heading 2"):
            level = 2
        elif style.startswith("heading 3"):
            level = 3
        elif _CN_SEQ_RE.match(text) and len(text) <= 40:
            level = 1
        if level and text not in seen:
            seen.add(text)
            out.append({"level": level, "title": text})
    return out


def load_bid_template(project_id) -> dict:
    """读取模板元信息：name/size/path/structure；无模板返回 None"""
    tdir = _template_dir(project_id)
    if not os.path.isdir(tdir):
        return None
    for fn in sorted(os.listdir(tdir)):
        if fn.lower().endswith(".docx"):
            p = os.path.join(tdir, fn)
            return {
                "name": fn,
                "size": os.path.getsize(p),
                "path": p,
                "structure": parse_template_structure(p),
            }
    return None


def save_bid_template(project_id, file_bytes: bytes, filename: str) -> dict:
    """保存投标模板（单模板覆盖，仅接受 .docx）；返回模板元信息"""
    if not (filename or "").lower().endswith(".docx"):
        raise ValueError("仅支持上传 .docx 模板")
    tdir = _template_dir(project_id)
    os.makedirs(tdir, exist_ok=True)
    # 单模板覆盖：清空旧模板
    for fn in os.listdir(tdir):
        if fn.lower().endswith(".docx"):
            try:
                os.remove(os.path.join(tdir, fn))
            except Exception:
                pass
    name = _safe_name(filename)
    with open(os.path.join(tdir, name), "wb") as f:
        f.write(file_bytes or b"")
    return load_bid_template(project_id)


def delete_bid_template(project_id) -> bool:
    """删除模板；返回是否存在模板"""
    tdir = _template_dir(project_id)
    if not os.path.isdir(tdir):
        return False
    deleted = False
    for fn in os.listdir(tdir):
        if fn.lower().endswith(".docx"):
            try:
                os.remove(os.path.join(tdir, fn))
                deleted = True
            except Exception:
                pass
    return deleted


def _md_sections(lines: list) -> list:
    """md 行 → 章节 [{level, title, body:[...]}]"""
    secs, cur = [], None
    for ln in lines:
        ln = ln.rstrip()
        m = re.match(r"^(#{1,3})\s+(.*)", ln)
        if m:
            cur = {"level": len(m.group(1)), "title": m.group(2).strip(), "body": []}
            secs.append(cur)
        elif cur is not None and ln.strip() and ln.strip() != "---":
            cur["body"].append(ln)
    return secs


def _render_docx_with_template(md_path: str, out_path: str, proj: dict,
                               template_path: str, spec: dict) -> None:
    """模板化导出：打开模板副本，按章节标题匹配插入生成内容，复用模板样式（FR-14）"""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    document = Document(template_path)
    sections = _md_sections(open(md_path, encoding="utf-8").read().splitlines())
    paragraphs = list(document.paragraphs)

    def _insert_after(anchor, text: str, style: str = None):
        new_p = OxmlElement("w:p")
        anchor.addnext(new_p)
        np = Paragraph(new_p, anchor.getparent())
        if text:
            run = np.add_run(text)
        if style:
            try:
                np.style = style
            except Exception:
                pass
        return new_p

    for sec in sections:
        norm = _norm_title(sec["title"])
        target = None
        # 精确归一化匹配优先，其次双向包含匹配
        for p in paragraphs:
            if norm and _norm_title(p.text) == norm:
                target = p
                break
        if target is None:
            for p in paragraphs:
                t = _norm_title(p.text)
                if norm and t and (norm in t or t in norm):
                    target = p
                    break
        if target is None:
            # 未匹配 → 追加文末
            document.add_heading(sec["title"], level=min(sec["level"], 3))
            for ln in sec["body"]:
                if ln.startswith("- ") or ln.startswith("* "):
                    document.add_paragraph(_md_inline(ln[2:]), style="List Bullet")
                elif ln.startswith("> "):
                    document.add_paragraph(_md_inline(ln[2:]), style="Quote")
                else:
                    document.add_paragraph(_md_inline(ln))
        else:
            anchor = target._p
            for ln in sec["body"]:
                bullet = ln.startswith("- ") or ln.startswith("* ")
                text = _md_inline(ln[2:] if bullet else ln)
                anchor = _insert_after(anchor, text, "List Bullet" if bullet else None)
    document.save(out_path)


# ────────────────────────────────────────────
# 技术方案演示网页（FR-15）
# ────────────────────────────────────────────

DEMO_SKELETON = os.path.join(config.BASE_DIR, "static", "demo-preview.html")


def _esc(s) -> str:
    """HTML 转义（插入模板前防注入）"""
    return (str(s or "").replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))


def _demo_features(report: dict, technical_text: str) -> list:
    """功能清单：tech_params + 规范书 LLM 提炼合并去重（FR-15）"""
    features, seen = [], set()
    for t in report.get("tech_params") or []:
        item = (t.get("item") or "").strip()
        if not item or item == "未识别":
            continue
        key = _norm_title(item)
        if key in seen:
            continue
        seen.add(key)
        features.append({
            "name": item,
            "desc": t.get("value") or "",
            "source": t.get("source") or "拆标报告",
            "is_key": bool(t.get("is_key")),
            "type": "tech",
        })
    # LLM 提炼补充（失败静默降级，不阻塞生成）
    if (technical_text or "").strip():
        sys_p = ("你是运维产品专家。请从招标技术要求中提炼可展示的系统功能点，"
                 "输出严格 JSON（不要输出其他内容）："
                 '{"features": [{"name": "功能名", "desc": "一句话说明"}]}，3~8 条。')
        got = _call_llm_json(sys_p, technical_text[:LLM_CHUNK_SIZE],
                             trace_ctx={"stage": "bid_demo_features"})
        for f in (got or {}).get("features") or []:
            if not isinstance(f, dict) or not f.get("name"):
                continue
            key = _norm_title(f["name"])
            if key in seen:
                continue
            seen.add(key)
            features.append({
                "name": f["name"],
                "desc": f.get("desc") or "",
                "source": "规范书技术要求",
                "is_key": False,
                "type": "llm",
            })
    return features


def _demo_feature_block(features: list) -> str:
    """功能响应清单 HTML 表格（数据由拆标报告驱动）"""
    if not features:
        return ""
    rows = []
    for i, f in enumerate(features[:12], 1):
        star = "★ " if f.get("is_key") else ""
        rows.append(
            f'<tr><td>{i}</td><td>{star}{_esc(f.get("name"))}</td>'
            f'<td>{_esc(f.get("desc"))}</td><td>{_esc(f.get("source"))}</td></tr>'
        )
    return (
        '<div class="panel" style="margin-bottom:16px">'
        '<h3>技术参数响应清单<span class="tag">拆标报告驱动</span></h3>'
        '<div class="meta">来源：拆标报告技术参数与规范书技术要求 · ★ 为关键参数</div>'
        '<table><thead><tr><th style="width:40px">#</th><th>参数项</th>'
        '<th>要求值 / 说明</th><th style="width:200px">来源</th></tr></thead>'
        f"<tbody>{''.join(rows)}</tbody></table></div>"
    )


def _build_demo_html(proj: dict, features: list, params: list) -> str:
    """以 static/demo-preview.html 为蓝本，数据驱动生成自包含单 HTML（FR-15）"""
    try:
        with open(DEMO_SKELETON, encoding="utf-8") as f:
            html = f.read()
    except Exception:
        html = "<html><head><meta charset='utf-8'><title>投标演示</title></head><body><h1>演示网页骨架缺失</h1></body></html>"
    name = proj.get("name") or "NeuOps 智慧运维平台"
    tenderee = proj.get("tenderee") or "招标方"
    date = time.strftime("%Y-%m-%d")
    html = html.replace(
        "NeuOps 智慧运维平台 · 资源总览大屏（投标演示原型）",
        f"{name} · 资源总览大屏（投标演示）",
    )
    html = html.replace(
        "本页为投标演示「模拟界面原型」，数据由拆标报告技术参数驱动，仅用于效果展示。",
        f"本页为投标演示「模拟界面原型」，数据由拆标报告技术参数驱动（{tenderee} · {date}）。",
    )
    # 插入技术参数响应清单（拆标报告驱动）
    block = _demo_feature_block(features)
    if block:
        html = html.replace('<div class="footer">', block + '<div class="footer">', 1)
    return html


# ────────────────────────────────────────────
# 导出
# ────────────────────────────────────────────

def export_document(project_id: int, doc_id: str, fmt: str = "md"):
    """导出成果：md/html 直接返回路径；docx 渲染后返回路径（有模板走模板化导出，FR-14）"""
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    doc = next((d for d in (proj.get("generated_docs") or []) if d.get("id") == doc_id), None)
    if not doc:
        raise ValueError("成果不存在")
    src_path = doc.get("path", "")
    if not os.path.isfile(src_path):
        raise ValueError("成果文件缺失")
    doc_type = doc.get("type", "")
    if fmt == "md":
        db.bid_set_status(project_id, "已导出")
        return src_path
    if fmt == "html":
        if doc_type == "tech_demo" and src_path.endswith(".html"):
            db.bid_set_status(project_id, "已导出")
            return src_path
        raise ValueError("该成果不支持 html 导出")
    if fmt == "docx":
        out_dir = os.path.dirname(src_path)
        out_path = os.path.join(out_dir, f"{doc_id}.docx")
        spec = _merge_format_spec(load_format_spec(project_id))
        tpl = load_bid_template(project_id)
        if tpl and doc_type != "tech_demo":
            _render_docx_with_template(src_path, out_path, proj, tpl["path"], spec)
        else:
            _render_docx(src_path, out_path, proj, spec, doc_type=doc_type)
        db.bid_set_status(project_id, "已导出")
        return out_path
    raise ValueError(f"不支持的导出格式: {fmt}")


def _render_docx(md_path: str, out_path: str, proj: dict, spec: dict, doc_type: str = "") -> None:
    """按排版基线渲染 docx：中文字体 / 标题 / 页边距 / 封面 / 目录 / 页码（FR-11/FR-12）
    技术方案导出时在"技术方案要点"后插入功能截图区 Word 表格骨架（FR-15）"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    document = Document()
    sec = document.sections[0]
    sec.top_margin = Cm(spec["margin_top_cm"])
    sec.bottom_margin = Cm(spec["margin_bottom_cm"])
    sec.left_margin = Cm(spec["margin_left_cm"])
    sec.right_margin = Cm(spec["margin_right_cm"])

    # 正文样式：宋体小四、1.5 倍行距、首行缩进（显式 eastAsia 中文字体回退）
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(spec["size_body"])
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), spec["font_body"])
    npf = normal.paragraph_format
    npf.line_spacing = spec["line_spacing"]
    npf.first_line_indent = Cm(spec["first_line_indent_cm"])
    npf.space_after = Pt(6)

    _style_heading(document, "Heading 1", spec["font_heading"], spec["size_h1"],
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_heading(document, "Heading 2", spec["font_heading"], spec["size_h2"])
    _style_heading(document, "Heading 3", spec["font_heading"], spec["size_h3"])

    lines = open(md_path, encoding="utf-8").read().splitlines()
    headings = [ln.strip() for ln in lines if re.match(r"^#{1,3}\s+\S", ln.strip())]

    if spec.get("cover"):
        _add_cover(document, proj, spec)
    if spec.get("toc") and headings:
        _add_toc(document, headings, spec)

    shot_inserted = False
    for ln in lines:
        ln = ln.rstrip()
        if not ln.strip():
            continue
        m = re.match(r"^(#{1,3})\s+(.*)", ln)
        if m:
            document.add_heading(m.group(2), level=len(m.group(1)))
            if (doc_type == "tech_proposal" and not shot_inserted
                    and _norm_title(m.group(2)) == "技术方案要点"):
                _add_demo_shot_table(document, proj, spec)
                shot_inserted = True
        elif ln.lstrip().startswith("!["):
            m_img = re.match(r"^!\[(.*?)\]\(([^)]+)\)$", ln.strip())
            if not m_img:
                document.add_paragraph(_md_inline(ln))
                continue
            img_rel = m_img.group(2)
            img_abs = os.path.join(os.path.dirname(md_path), img_rel)
            if os.path.isfile(img_abs):
                try:
                    document.add_picture(img_abs, width=Cm(14))
                except Exception:
                    document.add_paragraph(f"（截图插入失败：{img_rel}）")
            else:
                document.add_paragraph(f"（截图缺失：{img_rel}）")
            continue
        elif ln.startswith("- ") or ln.startswith("* "):
            document.add_paragraph(_md_inline(ln[2:]), style="List Bullet")
        elif ln.startswith("> "):
            document.add_paragraph(_md_inline(ln[2:]), style="Quote")
        elif ln.strip() == "---":
            continue
        else:
            document.add_paragraph(_md_inline(ln))

    if spec.get("page_number"):
        _add_page_number(sec)
    document.save(out_path)


def _add_demo_shot_table(document, proj: dict, spec: dict) -> None:
    """技术方案导出：在'技术方案要点'后插入功能截图区（FR-15）
    有 playwright 截图 → Word 表格真插图；无截图/失败 → 文字占位表格"""
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    p = document.add_paragraph()
    r = p.add_run("【功能截图区 · 模拟界面原型】")
    r.bold = True
    r.font.size = Pt(12)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), spec["font_heading"])
    shots = []
    try:
        from .screenshot import load_shots
        shots = load_shots(proj.get("id"))
    except Exception:
        shots = []
    if shots:
        rows = min(len(shots), 4)
        tbl = document.add_table(rows=rows + 1, cols=2)
        try:
            tbl.style = "Table Grid"
        except Exception:
            pass
        hdr = tbl.rows[0].cells
        hdr[0].text = "功能截图"
        hdr[1].text = "说明"
        for i, s in enumerate(shots[:rows], 1):
            cell0, cell1 = tbl.rows[i].cells
            cell1.text = s.get("name") or s.get("section") or "演示界面"
            fp = s.get("path") or ""
            if os.path.isfile(fp):
                try:
                    cell0.paragraphs[0].add_run().add_picture(fp, width=Cm(12.5))
                except Exception:
                    cell0.text = f"（截图插入失败：{s.get('file')}）"
            else:
                cell0.text = "（此处为系统界面截图占位，交付时插入高清截图）"
    else:
        tbl = document.add_table(rows=3, cols=2)
        try:
            tbl.style = "Table Grid"
        except Exception:
            pass
        hdr = tbl.rows[0].cells
        hdr[0].text = "功能截图"
        hdr[1].text = "说明"
        c1 = tbl.rows[1].cells
        c1[0].text = "（此处为系统界面截图占位，交付时插入高清截图）"
        c1[1].text = "资源总览大屏：KPI 指标、趋势折线图、告警分布环形图、主机状态、告警列表"
        c2 = tbl.rows[2].cells
        c2[0].text = "（对应演示网页：static/demo-preview.html）"
        c2[1].text = f"演示网页由拆标报告技术参数驱动生成（{proj.get('name', '')}）"


def _style_heading(document, style_name: str, font: str, size_pt, align=None) -> None:
    """标题样式：中文字体 + 黑色 + 字号（去除 python-docx 默认蓝色）"""
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    st = document.styles[style_name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size_pt)
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(6)
    if align is not None:
        st.paragraph_format.alignment = align


def _add_cover(document, proj: dict, spec: dict) -> None:
    """封面：项目名称 / 投标文件 / 招标方 / 日期"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt
    name = proj.get("name") or "投标项目"
    tenderee = proj.get("tenderee") or ""
    date = time.strftime("%Y年%m月%d日")
    for _ in range(5):
        document.add_paragraph("")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.font.size = Pt(22)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), spec["font_heading"])
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("投 标 文 件")
    r2.font.size = Pt(26)
    r2._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), spec["font_heading"])
    if tenderee:
        p3 = document.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run(f"招标方：{tenderee}").font.size = Pt(14)
    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(date).font.size = Pt(14)
    document.add_page_break()


def _add_toc(document, headings: list, spec: dict) -> None:
    """静态目录页（基于文档标题行）"""
    document.add_heading("目  录", level=1)
    for i, h in enumerate(headings, 1):
        clean = re.sub(r"^#{1,3}\s+", "", h).strip()
        document.add_paragraph(f"{i}. {clean}")
    document.add_page_break()


def _add_page_number(sec) -> None:
    """页脚居中页码（PAGE 域）"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def _md_inline(s: str) -> str:
    """md 行内简化：去除 ** 加粗标记（渲染层已统一样式）"""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s or "")


# ────────────────────────────────────────────
# 分步编写流水线（NO-009 FR-17/18/19/20/21）
# ────────────────────────────────────────────

def requirements_analysis(project_id: int) -> dict:
    """需求分析：LLM 把需求整理为开发可读的结构化 PRD；失败降级拆标摘要（FR-17）"""
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    text = _read_extracted_text(project_id)
    report = proj.get("parse_report") or {}
    report_md = _report_to_md(report)
    sys_p = ("你是资深需求分析师。请基于招标文件与拆标报告，把需求整理为开发能读懂的结构化 PRD，"
             "输出严格 JSON（不要输出其他内容），结构："
             '{"summary":"项目概述","roles":["用户角色"],'
             '"features":[{"name":"功能点","desc":"说明","priority":"P0/P1/P2"}],'
             '"pages":[{"name":"页面","desc":"说明","features":["关联功能点"]}],'
             '"interactions":[{"page":"页面","action":"用户操作","result":"系统响应"}],'
             '"acceptance":["验收口径"]}。'
             "不要编造招标文件中没有的内容，无法推断的用「待补充」标注。")
    user = f"【项目】{proj.get('name')}\n【拆标报告】\n{report_md[:4000]}\n\n【招标原文】\n{text[:6000]}"
    got = _call_llm_json(sys_p, user, trace_ctx={"stage": "bid_requirements"})
    if not got or not isinstance(got, dict) or not (got.get("features") or got.get("pages")):
        # LLM 未返回有效结构（含空结构 dict）→ 降级规则 PRD（FR-17）
        got = _fallback_prd(proj, report)
    # 补齐必填键，避免前端渲染缺字段
    for k in ("summary", "roles", "features", "pages", "interactions", "acceptance"):
        got.setdefault(k, [] if k != "summary" else "")
    db.bid_save_prd(project_id, got)
    return got


def _fallback_prd(proj: dict, report: dict) -> dict:
    """降级 PRD：从拆标报告生成最小结构（FR-17 降级）"""
    features = []
    for t in report.get("tech_params") or []:
        item = (t.get("item") or "").strip()
        if item and item != "未识别":
            features.append({"name": item, "desc": t.get("value") or "", "priority": "P1"})
    return {
        "summary": f"{proj.get('name')} 建设与服务需求（依据拆标报告整理）",
        "roles": ["系统管理员", "运维工程师", "业务用户"],
        "features": features[:20] or [{"name": "基础运维能力", "desc": "待补充", "priority": "P1"}],
        "pages": [{"name": "资源总览大屏", "desc": "KPI 与告警总览", "features": []}],
        "interactions": [],
        "acceptance": ["满足招标文件全部技术参数要求"],
    }


def generate_mockup(project_id: int) -> dict:
    """假页面生成：LLM 基于 PRD + 拆标参数产出演示 HTML；失败降级规则版本（FR-18/FR-15）"""
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    prd = proj.get("prd_json") or {}
    report = proj.get("parse_report") or {}
    if not prd:
        prd = _fallback_prd(proj, report)
    technical_text = _read_extracted_text(project_id, categories=("technical",))
    features = _demo_features(report, technical_text)
    body = _llm_mockup_html(proj, prd, report)
    source = "llm" if body else "rule"
    if not body:
        body = _build_demo_html(proj, features, report.get("tech_params") or [])
    output_dir = os.path.join(BID_UPLOAD_ROOT, str(project_id), "outputs")
    os.makedirs(output_dir, exist_ok=True)
    doc_id = uuid.uuid4().hex[:8]
    path = os.path.join(output_dir, f"{doc_id}.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(body)
    doc = {
        "id": doc_id,
        "type": "tech_demo",
        "title": "演示网页（假页面）",
        "path": path,
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "kind": "mockup",
        "source": source,  # llm：LLM 直出完整 HTML；rule：规则版兜底（FR-18）
    }
    db.bid_add_generated_doc(project_id, doc)
    db.bid_set_status(project_id, "已生成")
    return {"doc": doc, "html_preview": body[:200]}


def _llm_mockup_html(proj: dict, prd: dict, report: dict) -> str:
    """LLM 直出完整自包含 HTML 演示页（FR-18）：
    - 侧边导航 + 顶栏 + Dashboard KPI 卡 + 内联 SVG 图表（折线/环形）+ 表格 + 状态徽标
    - 零外部依赖，可 file:// 直接打开；内容由 PRD + 拆标技术参数驱动
    - LLM 不可用/输出非完整 HTML → 返回 ""，由规则版 _build_demo_html 兜底
    """
    pages = prd.get("pages") or []
    features = prd.get("features") or []
    sys_p = (
        "你是资深前端设计工程师。请为投标演示生成**完整自包含的单文件 HTML 页面**（可直接 file:// 打开，零外部依赖）。\n"
        "硬性要求：\n"
        "1. 直接输出完整 HTML（以 <!DOCTYPE html> 开头、</html> 结尾），不要输出任何解释文字或代码块标记。\n"
        "2. 全部 CSS/JS 内嵌在 <style>/<script> 内，禁止外链 CDN、字体、图标库。\n"
        "3. 采用现代运维中台风格：深蓝侧边导航（含 logo 与菜单项，可点击跳锚点）+ 顶部栏（面包屑/搜索框/用户头像）+ 主内容区。\n"
        "4. 每个页面含 Dashboard KPI 卡片区，并使用**内联 SVG** 绘制至少 2 类图表：趋势折线图、告警分布环形图、资源占比柱状图（禁止图片外链）。\n"
        "5. 使用表格与状态徽标（运行正常/告警/离线等彩色圆点）。\n"
        "6. 设计规范：主色 #1e3a8a（政务蓝）、辅助 #3b82f6 / #10b981 / #f59e0b，背景 #f1f5f9，卡片白底圆角阴影，中文字体栈。\n"
        "7. 内容必须基于给定【PRD】与【技术参数】，可补充通用运维界面元素（KPI、告警列表、主机列表等），但不得编造具体数值承诺。\n"
        "8. 页面数量 3~6 个，至少包含「资源总览大屏」与「技术参数响应清单」两页；每页用 <section id=\"page-N\"> 包裹便于截图；参数清单用表格逐项列出（★ 标关键参数）。\n"
        "9. 浏览器视口 1440×900，内容区宽度约 1200px，页面顶部标题含项目名称。\n"
    )
    user = (f"【项目】{proj.get('name')}（招标方：{proj.get('tenderee') or '未填'}）\n"
            f"【PRD】\n{json.dumps(prd, ensure_ascii=False)[:4000]}\n"
            f"【技术参数】\n{json.dumps(report.get('tech_params') or [], ensure_ascii=False)[:2500]}")
    body = _call_llm_text(sys_p, user, trace_ctx={"stage": "bid_mockup_html"})
    if not body or "<html" not in body.lower() or "</html>" not in body.lower():
        return ""
    if "section" not in body.lower() or "svg" not in body.lower():
        return ""
    return body


def _render_mockup_html(proj: dict, prd: dict, page_data: list, features: list) -> str:
    """把 LLM 页面骨架渲染为自包含单 HTML（零依赖，导航 + 各页 + 参数响应清单）"""
    name = proj.get("name") or "NeuOps 智慧运维平台"
    tenderee = proj.get("tenderee") or "招标方"
    date = time.strftime("%Y-%m-%d")
    nav_items = "".join(
        f'<a href="#page-{i}" class="nav-item">{_esc(p.get("name") or f"页面{i+1}")}</a>'
        for i, p in enumerate(page_data)
    )
    sections = []
    for i, p in enumerate(page_data):
        kpis = ""
        for k in (p.get("kpis") or [])[:6]:
            kpis += (f'<div class="kpi"><div class="kpi-name">{_esc(k.get("name"))}</div>'
                     f'<div class="kpi-value">{_esc(k.get("value"))}'
                     f'<span class="kpi-trend">{_esc(k.get("trend") or "")}</span></div></div>')
        blocks = ""
        for b in (p.get("blocks") or [])[:6]:
            items = "".join(f"<li>{_esc(x)}</li>" for x in (b.get("items") or [])[:8])
            blocks += (f'<div class="panel"><h3>{_esc(b.get("title"))}</h3>'
                       f'<ul class="block-items">{items}</ul></div>')
        sections.append(
            f'<section id="page-{i}" class="page"><h2>{_esc(p.get("name"))}</h2>'
            f'<p class="page-desc">{_esc(p.get("desc"))}</p>'
            f'<div class="kpi-grid">{kpis}</div><div class="panel-grid">{blocks}</div></section>'
        )
    param_rows = ""
    for i, f in enumerate(features[:12], 1):
        star = "★ " if f.get("is_key") else ""
        param_rows += (f'<tr><td>{i}</td><td>{star}{_esc(f.get("name"))}</td>'
                       f'<td>{_esc(f.get("desc"))}</td><td>{_esc(f.get("source"))}</td></tr>')
    params_block = (f'<section id="params" class="page"><h2>技术参数响应清单</h2>'
                    f'<div class="panel"><div class="meta">来源：拆标报告技术参数 · ★ 为关键参数</div>'
                    f'<table><thead><tr><th style="width:40px">#</th><th>参数项</th><th>要求值 / 说明</th>'
                    f'<th style="width:200px">来源</th></tr></thead><tbody>{param_rows}</tbody></table>'
                    f'</div></section>')
    return f"""<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{_esc(name)} · 投标演示原型</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;background:#f0f2f5;color:#1f2937}}
.nav{{position:sticky;top:0;background:#0f172a;color:#fff;display:flex;gap:4px;padding:10px 20px;overflow-x:auto;z-index:10}}
.nav-item{{color:#cbd5e1;text-decoration:none;padding:6px 14px;border-radius:6px;font-size:14px;white-space:nowrap}}
.nav-item:hover,.nav-item.active{{background:#1e3a8a;color:#fff}}
.hero{{background:linear-gradient(135deg,#1e3a8a,#3b82f6);color:#fff;padding:36px 24px;text-align:center}}
.hero h1{{font-size:24px;margin-bottom:8px}}.hero p{{opacity:.85;font-size:14px}}
main{{max-width:1200px;margin:0 auto;padding:20px}}
.page{{margin-bottom:24px}}
.page h2{{font-size:18px;border-left:4px solid #2563eb;padding-left:10px;margin:16px 0 12px}}
.page-desc{{color:#6b7280;font-size:13px;margin-bottom:12px}}
.kpi-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:12px;margin-bottom:16px}}
.kpi{{background:#fff;border-radius:10px;padding:14px;box-shadow:0 1px 3px rgba(0,0,0,.08)}}
.kpi-name{{font-size:12px;color:#6b7280}}.kpi-value{{font-size:20px;font-weight:600;margin-top:4px}}
.kpi-trend{{font-size:12px;color:#16a34a;margin-left:6px}}
.panel-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(260px,1fr));gap:12px}}
.panel{{background:#fff;border-radius:10px;padding:14px;box-shadow:0 1px 3px rgba(0,0,0,.08);margin-bottom:12px}}
.panel h3{{font-size:14px;margin-bottom:8px;color:#111827}}
.block-items{{list-style:none;font-size:13px;color:#374151;line-height:1.8}}
.block-items li::before{{content:"· ";color:#2563eb}}
table{{width:100%;border-collapse:collapse;background:#fff;font-size:13px}}
th,td{{border:1px solid #e5e7eb;padding:8px 10px;text-align:left}}
th{{background:#f8fafc;color:#111827}}
.meta{{font-size:12px;color:#6b7280;margin-bottom:8px}}
.tag{{display:inline-block;background:#dbeafe;color:#1e40af;font-size:11px;padding:2px 8px;border-radius:999px;margin-left:8px}}
.footer{{text-align:center;color:#9ca3af;font-size:12px;padding:24px;border-top:1px solid #e5e7eb}}
</style></head>
<body>
<div class="nav">{nav_items}<a href="#params" class="nav-item">参数清单</a></div>
<div class="hero"><h1>{_esc(name)}</h1><p>投标演示「模拟界面原型」· {_esc(tenderee)} · {date}</p></div>
<main>{''.join(sections)}{params_block}</main>
<div class="footer">本页为投标演示「模拟界面原型」，数据由需求分析 PRD 与拆标报告驱动，仅用于效果展示。</div>
<script>document.querySelectorAll('.nav-item').forEach(a=>a.addEventListener('click',()=>{{
document.querySelectorAll('.nav-item').forEach(x=>x.classList.remove('active'));a.classList.add('active');}}));</script>
</body></html>"""


def generate_outline(project_id: int) -> list:
    """章节大纲：LLM 基于模板章节树 + 拆标报告生成；失败降级模板章节（FR-19）"""
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    report = proj.get("parse_report") or {}
    report_md = _report_to_md(report)
    tpl = load_bid_template(project_id)
    tpl_titles = [s.get("title") for s in (tpl or {}).get("structure") or [] if s.get("title")]
    sys_p = ("你是招投标方案专家。请基于投标模板章节结构与拆标报告，输出技术方案建议书章节大纲，"
             "输出严格 JSON 数组（不要输出其他内容）："
             '[{"title":"章节标题","purpose":"本章应覆盖的要点"}]，8~16 章。')
    user = (f"【项目】{proj.get('name')}\n【模板章节】\n{json.dumps(tpl_titles, ensure_ascii=False)}\n"
            f"【拆标报告】\n{report_md[:4000]}")
    got = _call_llm_json(sys_p, user, trace_ctx={"stage": "bid_outline"})
    items = got if isinstance(got, list) else (got or {}).get("chapters")
    if not items or not isinstance(items, list):
        items = [{"title": t or f"章节{i+1}", "purpose": ""} for i, t in enumerate(tpl_titles)] or \
                _fallback_outline()
    outline = []
    for i, it in enumerate(items[:20], 1):
        if isinstance(it, dict) and it.get("title"):
            outline.append({"index": i, "title": it["title"], "purpose": it.get("purpose") or ""})
    if not outline:
        outline = _fallback_outline()
    db.bid_save_outline(project_id, outline)
    chapters = [{"index": c["index"], "title": c["title"], "purpose": c.get("purpose", ""),
                 "content": "", "confirmed": False, "source": "",
                 "source_text": _match_source_text(c["title"], report)} for c in outline]
    db.bid_save_chapters(project_id, chapters)
    return outline


def _match_source_text(title: str, report: dict) -> str:
    """根据章节标题匹配拆标报告对应节，生成左侧「规范书要求」依据（FR-20）"""
    text = _report_to_md(report)
    if not text:
        return ""
    if any(k in title for k in ("参数", "响应", "性能", "功能")):
        sec = "## 技术参数"
    elif any(k in title for k in ("资质", "资格")):
        sec = "## 资质要求"
    elif any(k in title for k in ("业绩", "案例")):
        sec = "## 业绩要求"
    elif any(k in title for k in ("风险", "应对")):
        sec = "## 废标/否决条款"
    elif any(k in title for k in ("实施", "进度", "计划")):
        sec = "## 评分标准"
    else:
        sec = ""
    if not sec:
        return text[:1500]
    # 截取该节内容
    idx = text.find(sec)
    if idx < 0:
        return text[:1500]
    nxt = text.find("\n## ", idx + 2)
    return text[idx:(nxt if nxt > 0 else len(text))][:2500]


def _fallback_outline() -> list:
    """降级大纲：标准技术方案章节（FR-19 降级）"""
    titles = [("项目概述", "项目背景、建设目标、建设内容"), ("总体技术方案", "系统架构、技术路线、部署架构"),
              ("功能设计", "各功能模块详细设计"), ("技术参数响应", "逐项响应拆标技术参数"),
              ("实施方案", "实施计划、里程碑、资源投入"), ("运维与服务方案", "运维体系、SLA、服务承诺"),
              ("项目团队与资质", "团队配置、资质证书"), ("风险分析与应对", "风险识别与应对措施"),
              ("培训与售后", "培训计划、售后服务体系"), ("服务承诺与验收", "验收标准、服务承诺")]
    return [{"index": i, "title": t, "purpose": p} for i, (t, p) in enumerate(titles, 1)]


def generate_chapter(project_id: int, index: int, force: bool = False) -> dict:
    """逐章生成：单章一次 LLM 调用；失败标记「待补充」不中断（FR-19/FR-20）"""
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    chapters = proj.get("chapters_json") or []
    if not chapters:
        raise ValueError("请先生成章节大纲")
    if index < 1 or index > len(chapters):
        raise ValueError("章节序号无效")
    ch = chapters[index - 1]
    if ch.get("confirmed") and not force:
        return {"chapter": ch, "done": True}
    report = proj.get("parse_report") or {}
    report_md = _report_to_md(report)
    kb_ids = get_project_kb_ids(project_id)
    rag_lines = []
    for q in (ch.get("title"), ch.get("purpose") or "", "技术方案"):
        for hit in knowledge.search_knowledge(q, kb_ids, top_k=2):
            t = hit.get("summary") or hit.get("title") or ""
            if t and t not in rag_lines:
                rag_lines.append(t)
    rag_text = "\n".join(rag_lines[:12]) or "（知识库暂无相关内容）"
    confirmed_summary = _confirmed_summary(chapters, index)
    sys_p = (f"你是资深投标技术方案撰写专家。请撰写「{ch.get('title')}」章节正文，"
             "输出 Markdown（不要输出代码块标记），要求：内容具体、可直接用于投标文件；"
             "基于拆标报告与知识库依据，不编造未提供的数据；无法确认的项标注【待补充】。")
    user = (f"【章节】{ch.get('title')}\n【编写要点】{ch.get('purpose') or '（无）'}\n"
            f"【拆标报告】\n{report_md[:3500]}\n【知识库参考】\n{rag_text}\n"
            f"【已确认章节摘要】\n{confirmed_summary or '（无）'}")
    content = _call_llm_text(sys_p, user, trace_ctx={"stage": "bid_chapter", "chapter": index})
    if not content:
        content = _fallback_chapter(ch, report)
        ch["source"] = "rule"
    else:
        ch["source"] = "llm"
    ch["content"] = content
    db.bid_save_chapters(project_id, chapters)
    return {"chapter": ch, "done": False}


def _confirmed_summary(chapters: list, exclude_index: int) -> str:
    """已确认章节的内容摘要（供后续章节引用，避免重复与遗漏）"""
    lines = []
    for c in chapters:
        if c.get("index") == exclude_index or not c.get("confirmed"):
            continue
        title = c.get("title") or ""
        content = (c.get("content") or "").strip().replace("\n", " ")[:150]
        lines.append(f"- {title}：{content}")
    return "\n".join(lines)


def _fallback_chapter(ch: dict, report: dict) -> str:
    """降级章节：规则模板基于拆标报告生成（FR-19 降级，标注待补充）"""
    title = ch.get("title") or "本章节"
    purpose = ch.get("purpose") or ""
    lines = [f"## {purpose or '编写要点'}", ""]
    norm = _norm_title(title)
    if norm == "技术参数响应" or "参数" in title:
        lines.append("以下逐项响应拆标报告技术参数：")
        for t in report.get("tech_params") or []:
            item = t.get("item") or "未识别"
            if item == "未识别":
                continue
            lines.append(f"- {item}：要求 {t.get('value', '未识别')} → 我方应答：【待补充材料】")
    elif "资质" in title:
        lines.append("以下逐项响应资质要求：")
        for q in report.get("qualifications") or []:
            item = q.get("item") or "未识别"
            if item == "未识别":
                continue
            lines.append(f"- {item}（{q.get('level', '待补充')}）→ 我方资质：【待补充材料】")
    elif "业绩" in title:
        lines.append("以下逐项响应业绩要求：")
        for p in report.get("performance") or []:
            item = p.get("item") or "未识别"
            if item == "未识别":
                continue
            lines.append(f"- {item}（{p.get('requirement', '待补充')}）→ 我方业绩：【待补充材料】")
    else:
        lines.append(f"【待补充材料】本章「{title}」需依据拆标报告与项目实际资料补充编写。")
    lines += ["", f"> 本章由投标工作台按规则模板生成（LLM 不可用），请人工复核后使用。"]
    return "\n".join(lines)


def assemble_document(project_id: int) -> dict:
    """组装：合并全部已确认章节 → 最终文档，落盘 outputs/ 并进入成果列表（FR-21）"""
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    chapters = proj.get("chapters_json") or []
    confirmed = [c for c in chapters if c.get("confirmed") and (c.get("content") or "").strip()]
    if not confirmed:
        raise ValueError("尚未确认任何章节，无法组装")
    name = proj.get("name", "本项目")
    tenderee = proj.get("tenderee", "招标方")
    lines = [f"# {name} — 技术方案建议书", f"招标方：{tenderee}", ""]
    for c in confirmed:
        lines += [f"## {c.get('title')}", "", c.get("content") or "", ""]
    # 演示界面截图（FR-15）：引用 mockup-shots 相对路径，docx 导出时 _render_docx 真插入
    try:
        from .screenshot import SHOT_DIR_REL, load_shots
        _shots = load_shots(project_id)
    except Exception:
        _shots = []
    if _shots:
        lines += ["## 系统演示界面截图", ""]
        for s in _shots[:5]:
            lines.append(f"![{s.get('name') or s.get('section') or '演示界面'}]({SHOT_DIR_REL}/{s['file']})")
        lines += ["", "> 截图由 playwright 渲染演示网页生成（无浏览器环境自动降级为文字说明）。"]
    lines += ["---", "", "> 本文档由投标工作台分步生成，经逐章人工确认后组装。"]
    body = "\n".join(lines)
    doc_id = uuid.uuid4().hex[:8]
    output_dir = os.path.join(BID_UPLOAD_ROOT, str(project_id), "outputs")
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"{doc_id}.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write(body)
    doc = {
        "id": doc_id,
        "type": "tech_proposal",
        "title": "技术方案建议书（分步组装）",
        "path": path,
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "kind": "assemble",
    }
    db.bid_add_generated_doc(project_id, doc)
    db.bid_set_status(project_id, "已生成")
    return doc


# ────────────────────────────────────────────
# 一键智能起草流水线（NO-009 FR-22 / FR-23）
# ────────────────────────────────────────────

# 流水线阶段（与前端横向步骤条一一对应）
PIPELINE_STAGES = ["parse", "requirements", "mockup", "outline", "chapters", "shots", "assemble", "review"]
PIPELINE_STAGE_NAMES = {
    "parse": "拆标解析",
    "requirements": "需求分析",
    "mockup": "演示原型",
    "outline": "章节大纲",
    "chapters": "逐章编写",
    "shots": "界面截图",
    "assemble": "文档组装",
    "review": "人工复核",
}

# 内存进度表 {project_id: {...}}（进程内有效，按 pid 幂等覆盖；重启后前端重新触发）
_PIPELINE_PROGRESS: dict = {}


def _set_progress(project_id, stage, message="", status="running", chapter=None, chapter_total=0):
    """记录流水线进度快照（供前端步骤条轮询，FR-22）"""
    _PIPELINE_PROGRESS[project_id] = {
        "project_id": project_id,
        "stage": stage,
        "stage_index": PIPELINE_STAGES.index(stage) if stage in PIPELINE_STAGES else 0,
        "total": len(PIPELINE_STAGES),
        "chapter": chapter,
        "chapter_total": chapter_total,
        "message": message,
        "status": status,  # running / awaiting_review / completed / error
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }


def get_pipeline_status(project_id):
    """读取一键起草进度（前端轮询，FR-22）；未运行返回 idle 态"""
    st = _PIPELINE_PROGRESS.get(project_id)
    if not st:
        return {
            "project_id": project_id, "running": False, "stage": None,
            "stage_index": 0, "total": len(PIPELINE_STAGES),
            "chapter": None, "chapter_total": 0, "message": "",
            "status": "idle", "done": False, "updated_at": "",
        }
    return dict(st, running=st.get("status") == "running",
                done=st.get("status") in ("completed", "error"))


def _project_has_upload(project_id) -> bool:
    """项目目录下是否存在已上传规范书（排除 outputs/extracted/kb_docs/template 等子目录与元数据）"""
    d = os.path.join(BID_UPLOAD_ROOT, str(project_id))
    if not os.path.isdir(d):
        return False
    for fn in os.listdir(d):
        p = os.path.join(d, fn)
        if os.path.isfile(p) and fn not in ("format_spec.json", "project.json"):
            return True
    return False


def run_bid_pipeline(project_id: int, auto_confirm: bool = False) -> dict:
    """一键智能起草（FR-22）：
    校验已上传 → 拆标 → 需求分析 → 演示原型 → 章节大纲 → 逐章全量草稿 → 界面截图
    默认停在人工复核（awaiting_review），由页面逐章复核确认后手动组装；
    auto_confirm=True 全自动：确认全部章节 → 组装 → 合规自检 → 导出 docx（含真实截图）。
    任一步骤失败按既有降级规则不中断；截图失败降级文字占位（FR-15）。
    """
    proj = db.bid_get_project(project_id)
    if not proj:
        raise ValueError("项目不存在")
    if not _project_has_upload(project_id):
        raise ValueError("请先上传规范书再触发一键起草")
    current = "parse"
    try:
        _set_progress(project_id, "parse", "正在拆标解析…")
        parse_bid_document(project_id)
        current = "requirements"
        _set_progress(project_id, "requirements", "正在需求分析…")
        requirements_analysis(project_id)
        current = "mockup"
        _set_progress(project_id, "mockup", "正在生成演示原型…")
        generate_mockup(project_id)
        current = "outline"
        _set_progress(project_id, "outline", "正在生成章节大纲…")
        generate_outline(project_id)
        current = "chapters"
        chapters = (db.bid_get_project(project_id).get("chapters_json") or [])
        total = max(len(chapters), 1)
        for i, ch in enumerate(chapters, 1):
            _set_progress(project_id, "chapters",
                          f"正在编写第 {i}/{len(chapters)} 章：{ch.get('title') or ''}…",
                          chapter=i, chapter_total=total)
            try:
                generate_chapter(project_id, i, force=False)
            except Exception:
                continue  # 单章失败标「待补充」不中断（FR-20）
        current = "shots"
        _set_progress(project_id, "shots", "正在生成界面截图…")
        try:
            from .screenshot import shot_mockup
            shot_mockup(project_id)
        except Exception:
            pass  # 无浏览器/失败 → 降级文字占位，不阻断（FR-15）
        if not auto_confirm:
            _set_progress(project_id, "review", "草稿已全部生成，请逐章复核确认", status="awaiting_review")
            return {"status": "awaiting_review",
                    "message": "一键起草完成：拆标/需求/演示/大纲/逐章/截图已生成，请逐章复核确认后组装。"}
        # 全自动：确认全部章节 → 组装 → 合规自检 → 导出 docx（含真实截图）
        current = "assemble"
        _set_progress(project_id, "assemble", "自动确认章节并组装文档…")
        chapters = (db.bid_get_project(project_id).get("chapters_json") or [])
        for ch in chapters:
            ch["confirmed"] = True
        db.bid_save_chapters(project_id, chapters)
        doc = assemble_document(project_id)
        _set_progress(project_id, "review", "正在合规自检…")
        check = check_compliance(project_id)
        out_path = export_document(project_id, doc["id"], "docx")
        _set_progress(project_id, "review", "一键起草全部完成", status="completed")
        return {"status": "completed", "doc": doc, "docx": out_path, "check": check}
    except Exception as e:
        _set_progress(project_id, current, f"流水线中断：{e}", status="error")
        raise
